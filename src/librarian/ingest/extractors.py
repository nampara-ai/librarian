"""Text extraction adapters."""

from __future__ import annotations

import asyncio
import hashlib
import importlib
import json
import multiprocessing
import queue
import shutil
import subprocess
import tempfile
import time
import uuid
from collections.abc import Sequence
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Literal, Protocol, cast

from librarian.application.transcripts import (
    TranscriptFormat,
    parse_transcript,
    render_transcript,
)
from librarian.observability import sanitize_error_message

IMAGE_EXTENSIONS = frozenset({".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"})
ARCHIVE_EXTENSIONS = frozenset({".zip", ".tar", ".tgz", ".gz", ".bz2", ".xz", ".7z", ".rar"})
ZIP_CONTAINER_EXTENSIONS = frozenset({".docx", ".epub", ".pptx", ".xlsx"})
ARCHIVE_SIGNATURES = (
    (b"PK\x03\x04", "zip"),
    (b"PK\x05\x06", "zip"),
    (b"PK\x07\x08", "zip"),
    (b"\x1f\x8b", "gzip"),
    (b"BZh", "bzip2"),
    (b"\xfd7zXZ\x00", "xz"),
    (b"7z\xbc\xaf\x27\x1c", "7z"),
    (b"Rar!\x1a\x07\x00", "rar"),
    (b"Rar!\x1a\x07\x01\x00", "rar"),
)
TAR_USTAR_OFFSET = 257
_ARCHIVE_SIGNATURE_SAMPLE_BYTES = 4096
_MAX_PDF_PAGE_MANIFEST_BYTES = 256 * 1024 * 1024
OcrCorrectionMode = Literal["always", "never", "low-confidence"]
OcrPreprocessMode = Literal["none", "grayscale", "threshold", "deskew"]

OCR_CORRECTION_PROMPT = """You are correcting OCR text extracted from a PDF page.
Preserve every detail, name, number, heading, and paragraph. Fix OCR recognition errors,
line-break artifacts, spacing, and obvious punctuation issues. Do not summarize, omit,
invent, or reorder content. Return only corrected Markdown-compatible text."""


class OcrCorrectionProvider(Protocol):
    """Minimal LLM provider surface needed for OCR correction."""

    @property
    def name(self) -> str: ...

    async def complete(
        self,
        *,
        system_prompt: str,
        user_prompt: str,
        model: str,
        max_tokens: int,
        temperature: float,
    ) -> str: ...


class ExtractionMetrics(Protocol):
    """Metrics sink for extraction adapters."""

    def record_ocr_page(
        self,
        *,
        source: str,
        status: str,
        duration_ms: float,
        corrected: bool = False,
    ) -> None: ...


@dataclass(frozen=True, slots=True)
class PdfPageExtraction:
    """One extracted PDF page."""

    page_number: int
    text: str
    source: str
    confidence: float | None = None
    corrected: bool = False
    error: str | None = None
    raw_text: str | None = None
    image_path: str | None = None
    warnings: tuple[str, ...] = ()
    attempts: int = 0
    duration_ms: float | None = None


@dataclass(frozen=True, slots=True)
class OcrTextResult:
    """Raw OCR text plus optional engine confidence."""

    text: str
    confidence: float | None = None
    image_path: str | None = None


class TextFamilyExtractor:
    """Extractor for UTF-8 text-like files."""

    supported_extensions = frozenset({".txt", ".md", ".csv", ".json"})

    def __init__(self, *, max_input_bytes: int = 100 * 1024 * 1024) -> None:
        self.max_input_bytes = max_input_bytes

    async def extract(self, path: Path) -> str:
        return await asyncio.to_thread(self._extract_sync, path)

    def _extract_sync(self, path: Path) -> str:
        _validate_input_size(path, self.max_input_bytes, "Text extraction input")
        _validate_text_like(path)
        text = _read_limited_text_file(
            path,
            max_bytes=self.max_input_bytes,
            label="Text extraction input",
        )
        if path.suffix.lower() == ".json":
            try:
                parsed = json.loads(text)
            except json.JSONDecodeError:
                return text
            return json.dumps(parsed, indent=2, ensure_ascii=False)
        return text


class TranscriptFileExtractor:
    """Extractor for timestamped transcript caption files."""

    supported_extensions = frozenset({".srt", ".vtt"})

    def __init__(self, *, max_input_bytes: int = 100 * 1024 * 1024) -> None:
        self.max_input_bytes = max_input_bytes

    async def extract(self, path: Path) -> str:
        return await asyncio.to_thread(self._extract_sync, path)

    def _extract_sync(self, path: Path) -> str:
        _validate_input_size(path, self.max_input_bytes, "Transcript extraction input")
        _validate_text_like(path)
        text = _read_limited_text_file(
            path,
            max_bytes=self.max_input_bytes,
            label="Transcript extraction input",
        )
        segments = parse_transcript(text)
        if not segments:
            raise ValueError(f"No timestamped transcript segments found: {path}")
        return render_transcript(
            segments,
            format=TranscriptFormat.MARKDOWN,
            merge_sentences=True,
        )


class DocxExtractor:
    """Extractor for DOCX files."""

    supported_extensions = frozenset({".docx"})

    def __init__(self, *, max_input_bytes: int = 100 * 1024 * 1024) -> None:
        self.max_input_bytes = max_input_bytes

    async def extract(self, path: Path) -> str:
        return await asyncio.to_thread(self._extract_sync, path)

    def _extract_sync(self, path: Path) -> str:
        _validate_input_size(path, self.max_input_bytes, "DOCX extraction input")
        from docx import Document

        doc = Document(str(path))
        parts = [
            *(_paragraph_text(paragraph) for paragraph in doc.paragraphs),
            *(_table_text(table) for table in doc.tables),
        ]
        for section in doc.sections:
            parts.extend(_paragraph_text(paragraph) for paragraph in section.header.paragraphs)
            parts.extend(_table_text(table) for table in section.header.tables)
            parts.extend(_paragraph_text(paragraph) for paragraph in section.footer.paragraphs)
            parts.extend(_table_text(table) for table in section.footer.tables)
        return "\n\n".join(part for part in parts if part.strip())


class PdfExtractor:
    """Extractor for text-bearing PDFs."""

    supported_extensions = frozenset({".pdf"})

    def __init__(
        self,
        *,
        ocr_language: str = "eng",
        ocr_timeout_seconds: int = 120,
        ocr_pdf_dpi: int = 200,
        ocr_pdf_max_pages: int = 1_000,
        ocr_preprocess_mode: OcrPreprocessMode = "none",
        ocr_threshold: int = 180,
        ocr_preserve_page_images: bool = False,
        ocr_correction_provider: OcrCorrectionProvider | None = None,
        ocr_correction_mode: OcrCorrectionMode = "always",
        ocr_correction_model: str = "mock-cleaner",
        ocr_low_confidence_threshold: float = 85.0,
        ocr_max_correction_response_chars: int = 2 * 1024 * 1024,
        ocr_page_concurrency: int = 2,
        ocr_fail_on_page_error: bool = True,
        max_input_bytes: int = 200 * 1024 * 1024,
        max_pages: int = 1_000,
        metrics: ExtractionMetrics | None = None,
    ) -> None:
        self.ocr_language = ocr_language
        self.ocr_timeout_seconds = ocr_timeout_seconds
        self.ocr_pdf_dpi = ocr_pdf_dpi
        self.ocr_pdf_max_pages = ocr_pdf_max_pages
        _validate_ocr_preprocess_config(ocr_preprocess_mode, threshold=ocr_threshold)
        self.ocr_preprocess_mode: OcrPreprocessMode = ocr_preprocess_mode
        self.ocr_threshold = ocr_threshold
        self.ocr_preserve_page_images = ocr_preserve_page_images
        self.ocr_correction_provider = ocr_correction_provider
        self.ocr_correction_mode = ocr_correction_mode
        self.ocr_correction_model = ocr_correction_model
        self.ocr_low_confidence_threshold = ocr_low_confidence_threshold
        self.ocr_max_correction_response_chars = ocr_max_correction_response_chars
        self.ocr_page_concurrency = ocr_page_concurrency
        self.ocr_fail_on_page_error = ocr_fail_on_page_error
        self.max_input_bytes = max_input_bytes
        self.max_pages = max_pages
        self.metrics = metrics
        self.last_metadata: dict[str, object] | None = None
        self.page_manifest_path: Path | None = None

    def set_page_manifest_path(self, path: Path | None) -> None:
        """Set an optional durable manifest path for page-level extraction state."""
        self.page_manifest_path = path

    async def extract(self, path: Path) -> str:
        _validate_input_size(path, self.max_input_bytes, "PDF extraction input")
        pages, page_count = await asyncio.to_thread(self._extract_embedded_pages, path)
        if page_count > self.max_pages:
            raise ValueError(
                f"PDF has {page_count} pages, exceeding configured limit {self.max_pages}: {path}"
            )

        empty_pages = [page.page_number for page in pages if not page.text.strip()]
        if len(empty_pages) > self.ocr_pdf_max_pages:
            page_list = ", ".join(str(page_number) for page_number in empty_pages)
            raise ValueError(
                f"PDF contains {len(empty_pages)} scanned/empty pages, exceeding OCR page "
                f"limit {self.ocr_pdf_max_pages}: {page_list}"
            )

        source_sha256 = await asyncio.to_thread(_file_sha256, path)
        extraction_config = self._extraction_config()
        manifest = await _load_pdf_page_manifest(self.page_manifest_path)
        reusable_pages = _reusable_manifest_pages(
            manifest,
            source_sha256=source_sha256,
            extraction_config=extraction_config,
            page_count=page_count,
        )
        previous_attempts = _manifest_page_attempts(manifest)
        page_results: list[PdfPageExtraction | None] = []
        for page in pages:
            reusable_page = reusable_pages.get(page.page_number)
            if reusable_page is not None:
                page_results.append(reusable_page)
            elif page.text.strip():
                page_results.append(page)
            else:
                page_results.append(
                    PdfPageExtraction(
                        page_number=page.page_number,
                        text="",
                        source="pending",
                        attempts=previous_attempts.get(page.page_number, 0),
                    )
                )
        manifest_lock = asyncio.Lock()

        async def write_manifest() -> None:
            async with manifest_lock:
                await _write_pdf_page_manifest(
                    self.page_manifest_path,
                    source_path=path,
                    source_sha256=source_sha256,
                    extraction_config=extraction_config,
                    page_count=page_count,
                    pages=page_results,
                )

        await write_manifest()
        semaphore = asyncio.Semaphore(max(1, self.ocr_page_concurrency))

        async def recover_page(page_number: int) -> None:
            existing = page_results[page_number - 1]
            if existing is not None and existing.text.strip() and existing.source == "ocr":
                return
            async with semaphore:
                page_start = time.perf_counter()
                try:
                    ocr_result_obj = await asyncio.to_thread(
                        _ocr_pdf_page,
                        path,
                        page_number=page_number,
                        language=self.ocr_language,
                        timeout_seconds=self.ocr_timeout_seconds,
                        dpi=self.ocr_pdf_dpi,
                        preprocess_mode=self.ocr_preprocess_mode,
                        threshold=self.ocr_threshold,
                        image_artifact_path=self._ocr_page_image_artifact_path(page_number),
                    )
                    ocr_result = _coerce_ocr_result(ocr_result_obj)
                    corrected = await self._correct_ocr_page(
                        page_number=page_number,
                        text=ocr_result.text,
                        confidence=ocr_result.confidence,
                    )
                    corrected_page = corrected != ocr_result.text
                    duration_ms = (time.perf_counter() - page_start) * 1000
                    page_results[page_number - 1] = PdfPageExtraction(
                        page_number=page_number,
                        text=corrected,
                        source="ocr",
                        confidence=ocr_result.confidence,
                        corrected=corrected_page,
                        raw_text=ocr_result.text,
                        image_path=ocr_result.image_path,
                        attempts=previous_attempts.get(page_number, 0) + 1,
                        duration_ms=duration_ms,
                        warnings=_ocr_page_warnings(
                            confidence=ocr_result.confidence,
                            low_confidence_threshold=self.ocr_low_confidence_threshold,
                        ),
                    )
                    self._record_ocr_page(
                        status="succeeded",
                        duration_ms=duration_ms,
                        corrected=corrected_page,
                    )
                    await write_manifest()
                except Exception as exc:
                    duration_ms = (time.perf_counter() - page_start) * 1000
                    error = sanitize_error_message(exc)
                    self._record_ocr_page(
                        status="failed",
                        duration_ms=duration_ms,
                    )
                    if self.ocr_fail_on_page_error:
                        page_results[page_number - 1] = PdfPageExtraction(
                            page_number=page_number,
                            text="",
                            source="ocr",
                            error=error,
                            warnings=("ocr-page-failed",),
                            attempts=previous_attempts.get(page_number, 0) + 1,
                            duration_ms=duration_ms,
                        )
                        await write_manifest()
                        raise RuntimeError(
                            f"Unable to OCR scanned PDF page {page_number}: {error}"
                        ) from exc
                    page_results[page_number - 1] = PdfPageExtraction(
                        page_number=page_number,
                        text="",
                        source="ocr",
                        error=error,
                        warnings=("ocr-page-failed",),
                        attempts=previous_attempts.get(page_number, 0) + 1,
                        duration_ms=duration_ms,
                    )
                    await write_manifest()

        await asyncio.gather(*(recover_page(page_number) for page_number in empty_pages))
        final_pages = [page for page in page_results if page is not None]
        if not any(page.text.strip() for page in final_pages):
            raise ValueError(f"No extractable content found in PDF: {path}")

        self.last_metadata = {
            "artifact_type": "pdf-page-extraction",
            "manifest_path": str(self.page_manifest_path) if self.page_manifest_path else None,
            "source_sha256": source_sha256,
            "extraction_config": extraction_config,
            "page_count": page_count,
            "pages": [
                {
                    "page_number": page.page_number,
                    "source": page.source,
                    "chars": len(page.text),
                    "confidence": page.confidence,
                    "corrected": page.corrected,
                    "status": _pdf_page_status(page),
                    "error": page.error,
                    "warnings": list(page.warnings),
                    "attempts": page.attempts,
                    "duration_ms": page.duration_ms,
                    "image_path": page.image_path,
                }
                for page in final_pages
            ],
        }
        return render_pdf_pages_markdown(path, final_pages)

    def _record_ocr_page(
        self,
        *,
        status: str,
        duration_ms: float,
        corrected: bool = False,
    ) -> None:
        if self.metrics is None:
            return
        self.metrics.record_ocr_page(
            source="pdf",
            status=status,
            duration_ms=duration_ms,
            corrected=corrected,
        )

    def _extraction_config(self) -> dict[str, object]:
        return {
            "ocr_language": self.ocr_language,
            "ocr_timeout_seconds": self.ocr_timeout_seconds,
            "ocr_pdf_dpi": self.ocr_pdf_dpi,
            "ocr_preprocess_mode": self.ocr_preprocess_mode,
            "ocr_threshold": self.ocr_threshold,
            "ocr_preserve_page_images": self.ocr_preserve_page_images,
            "ocr_correction_mode": self.ocr_correction_mode,
            "ocr_correction_model": self.ocr_correction_model,
            "ocr_low_confidence_threshold": self.ocr_low_confidence_threshold,
            "ocr_max_correction_response_chars": self.ocr_max_correction_response_chars,
        }

    def _extract_sync(self, path: Path) -> str:
        """Compatibility wrapper for callers that still use synchronous extraction."""
        _validate_input_size(path, self.max_input_bytes, "PDF extraction input")
        pages, page_count = self._extract_embedded_pages(path)
        if page_count > self.max_pages:
            raise ValueError(
                f"PDF has {page_count} pages, exceeding configured limit {self.max_pages}: {path}"
            )
        empty_pages = [page.page_number for page in pages if not page.text.strip()]
        if len(empty_pages) > self.ocr_pdf_max_pages:
            raise ValueError(
                f"PDF contains {len(empty_pages)} scanned/empty pages, exceeding OCR page "
                f"limit {self.ocr_pdf_max_pages}: {path}"
            )
        page_outputs = list(pages)
        for page_number in empty_pages:
            try:
                ocr_result = _coerce_ocr_result(
                    _ocr_pdf_page(
                        path,
                        page_number=page_number,
                        language=self.ocr_language,
                        timeout_seconds=self.ocr_timeout_seconds,
                        dpi=self.ocr_pdf_dpi,
                        preprocess_mode=self.ocr_preprocess_mode,
                        threshold=self.ocr_threshold,
                        image_artifact_path=None,
                    )
                )
            except (RuntimeError, ValueError) as exc:
                error = sanitize_error_message(exc)
                raise RuntimeError(
                    f"Unable to OCR scanned PDF page {page_number}: {error}"
                ) from exc
            page_outputs[page_number - 1] = PdfPageExtraction(
                page_number=page_number,
                text=ocr_result.text,
                source="ocr",
                confidence=ocr_result.confidence,
                raw_text=ocr_result.text,
                image_path=ocr_result.image_path,
                warnings=_ocr_page_warnings(
                    confidence=ocr_result.confidence,
                    low_confidence_threshold=self.ocr_low_confidence_threshold,
                ),
            )
        self.last_metadata = {
            "artifact_type": "pdf-page-extraction",
            "page_count": page_count,
            "pages": [
                {
                    "page_number": page.page_number,
                    "source": page.source,
                    "chars": len(page.text),
                    "confidence": page.confidence,
                    "corrected": page.corrected,
                    "status": _pdf_page_status(page),
                    "error": page.error,
                    "warnings": list(page.warnings),
                    "attempts": page.attempts,
                    "duration_ms": page.duration_ms,
                    "image_path": page.image_path,
                }
                for page in page_outputs
            ],
        }
        return render_pdf_pages_markdown(path, page_outputs)

    def _ocr_page_image_artifact_path(self, page_number: int) -> Path | None:
        if not self.ocr_preserve_page_images or self.page_manifest_path is None:
            return None
        return self.page_manifest_path.with_name(
            f"{self.page_manifest_path.stem}.page-{page_number:04d}.png"
        )

    def _extract_embedded_pages(self, path: Path) -> tuple[list[PdfPageExtraction], int]:
        try:
            pdfplumber = importlib.import_module("pdfplumber")
        except ImportError as exc:
            raise RuntimeError("PDF support requires installing the 'pdf' extra") from exc

        pages: list[PdfPageExtraction] = []
        pdf_module = cast(Any, pdfplumber)
        with pdf_module.open(path) as pdf:
            page_count = len(pdf.pages)
            for page_number, page in enumerate(pdf.pages, start=1):
                page_text = cast(str | None, page.extract_text())
                if page_text and page_text.strip():
                    pages.append(
                        PdfPageExtraction(
                            page_number=page_number,
                            text=page_text,
                            source="embedded",
                        )
                    )
                else:
                    pages.append(
                        PdfPageExtraction(
                            page_number=page_number,
                            text="",
                            source="empty",
                        )
                    )
        return pages, page_count

    async def _correct_ocr_page(
        self,
        *,
        page_number: int,
        text: str,
        confidence: float | None = None,
    ) -> str:
        if self.ocr_correction_mode == "never":
            return text
        if (
            self.ocr_correction_mode == "low-confidence"
            and (
                confidence is None
                or confidence >= self.ocr_low_confidence_threshold
            )
        ):
            return text
        provider = self.ocr_correction_provider
        if provider is None:
            raise RuntimeError("LLM OCR correction is enabled but no LLM provider is configured")
        corrected = await provider.complete(
            system_prompt=OCR_CORRECTION_PROMPT,
            user_prompt=(
                f"Correct OCR text from PDF page {page_number}. "
                "Return only corrected text.\n\n"
                f"{text}"
            ),
            model=self.ocr_correction_model,
            max_tokens=8192,
            temperature=0.0,
        )
        clean = str(corrected).strip()
        if len(clean) > self.ocr_max_correction_response_chars:
            raise ValueError(
                f"LLM OCR correction returned {len(clean)} characters for page {page_number}, "
                "exceeding configured limit "
                f"{self.ocr_max_correction_response_chars}"
            )
        if not clean:
            raise ValueError(f"LLM OCR correction returned empty text for page {page_number}")
        return clean


class ImageOcrExtractor:
    """Extractor for image files through Tesseract OCR."""

    supported_extensions = IMAGE_EXTENSIONS

    def __init__(
        self,
        *,
        language: str = "eng",
        timeout_seconds: int = 120,
        preprocess_mode: OcrPreprocessMode = "none",
        threshold: int = 180,
    ) -> None:
        self.language = language
        self.timeout_seconds = timeout_seconds
        _validate_ocr_preprocess_config(preprocess_mode, threshold=threshold)
        self.preprocess_mode: OcrPreprocessMode = preprocess_mode
        self.threshold = threshold

    async def extract(self, path: Path) -> str:
        return await asyncio.to_thread(
            _ocr_image,
            path,
            language=self.language,
            timeout_seconds=self.timeout_seconds,
            preprocess_mode=self.preprocess_mode,
            threshold=self.threshold,
        )


class MarkItDownExtractor:
    """Optional broad-format extractor using Microsoft's MarkItDown."""

    supported_extensions = frozenset(
        {
            ".epub",
            ".html",
            ".htm",
            ".msg",
            ".pptx",
            ".rtf",
            ".wav",
            ".mp3",
            ".xls",
            ".xlsx",
            ".xml",
        }
    )

    def __init__(
        self,
        *,
        max_input_bytes: int = 50 * 1024 * 1024,
        timeout_seconds: int = 120,
    ) -> None:
        self.max_input_bytes = max_input_bytes
        self.timeout_seconds = timeout_seconds

    async def extract(self, path: Path) -> str:
        return await asyncio.to_thread(self._extract_with_timeout, path)

    def _extract_with_timeout(self, path: Path) -> str:
        self._validate_input_size(path)
        result_queue: multiprocessing.Queue[tuple[str, str]] = multiprocessing.Queue(maxsize=1)
        process = multiprocessing.Process(
            target=_markitdown_worker,
            args=(str(path), self.max_input_bytes, result_queue),
            daemon=True,
        )
        process.start()
        process.join(self.timeout_seconds)
        if process.is_alive():
            process.terminate()
            process.join()
            raise TimeoutError(f"Broad format conversion timed out after {self.timeout_seconds}s")
        try:
            status, payload = result_queue.get_nowait()
        except queue.Empty as exc:
            raise RuntimeError("Broad format conversion failed without returning a result") from exc
        if status == "ok":
            return payload
        raise RuntimeError(payload)

    def _extract_sync(self, path: Path) -> str:
        self._validate_input_size(path)
        return _extract_markitdown_sync(path)

    def _validate_input_size(self, path: Path) -> None:
        _validate_markitdown_input_size(path, self.max_input_bytes)
        reject_disallowed_archive_signature(path)


class CompositeExtractor:
    """Route extraction by file extension."""

    def __init__(
        self,
        *,
        ocr_language: str = "eng",
        ocr_timeout_seconds: int = 120,
        ocr_pdf_dpi: int = 200,
        ocr_pdf_max_pages: int = 1_000,
        ocr_preprocess_mode: OcrPreprocessMode = "none",
        ocr_threshold: int = 180,
        ocr_preserve_page_images: bool = False,
        ocr_correction_provider: OcrCorrectionProvider | None = None,
        ocr_correction_mode: OcrCorrectionMode = "always",
        ocr_correction_model: str = "mock-cleaner",
        ocr_low_confidence_threshold: float = 85.0,
        ocr_max_correction_response_chars: int = 2 * 1024 * 1024,
        ocr_page_concurrency: int = 2,
        ocr_fail_on_page_error: bool = True,
        text_max_input_bytes: int = 100 * 1024 * 1024,
        docx_max_input_bytes: int = 100 * 1024 * 1024,
        pdf_max_input_bytes: int = 200 * 1024 * 1024,
        pdf_max_pages: int = 1_000,
        universal_max_input_bytes: int = 50 * 1024 * 1024,
        universal_timeout_seconds: int = 120,
        metrics: ExtractionMetrics | None = None,
    ) -> None:
        extractors = [
            TextFamilyExtractor(max_input_bytes=text_max_input_bytes),
            TranscriptFileExtractor(max_input_bytes=text_max_input_bytes),
            DocxExtractor(max_input_bytes=docx_max_input_bytes),
            PdfExtractor(
                ocr_language=ocr_language,
                ocr_timeout_seconds=ocr_timeout_seconds,
                ocr_pdf_dpi=ocr_pdf_dpi,
                ocr_pdf_max_pages=ocr_pdf_max_pages,
                ocr_preprocess_mode=ocr_preprocess_mode,
                ocr_threshold=ocr_threshold,
                ocr_preserve_page_images=ocr_preserve_page_images,
                ocr_correction_provider=ocr_correction_provider,
                ocr_correction_mode=ocr_correction_mode,
                ocr_correction_model=ocr_correction_model,
                ocr_low_confidence_threshold=ocr_low_confidence_threshold,
                ocr_max_correction_response_chars=ocr_max_correction_response_chars,
                ocr_page_concurrency=ocr_page_concurrency,
                ocr_fail_on_page_error=ocr_fail_on_page_error,
                max_input_bytes=pdf_max_input_bytes,
                max_pages=pdf_max_pages,
                metrics=metrics,
            ),
            ImageOcrExtractor(
                language=ocr_language,
                timeout_seconds=ocr_timeout_seconds,
                preprocess_mode=ocr_preprocess_mode,
                threshold=ocr_threshold,
            ),
            MarkItDownExtractor(
                max_input_bytes=universal_max_input_bytes,
                timeout_seconds=universal_timeout_seconds,
            ),
        ]
        self._extractors = {
            extension: extractor
            for extractor in extractors
            for extension in extractor.supported_extensions
        }
        self.supported_extensions = frozenset(self._extractors)
        self.last_metadata: dict[str, object] | None = None

    async def extract(self, path: Path) -> str:
        extension = path.suffix.lower()
        if extension in ARCHIVE_EXTENSIONS:
            raise ValueError(f"Archive inputs are not supported by default: {extension}")
        extractor = self._extractors.get(extension)
        if extractor is None:
            raise ValueError(f"Unsupported file extension: {extension}")
        text = await extractor.extract(path)
        metadata = getattr(extractor, "last_metadata", None)
        self.last_metadata = metadata if isinstance(metadata, dict) else None
        return text

    def set_page_manifest_path(self, path: Path | None) -> None:
        """Forward page manifest paths to extractors that support them."""
        for extractor in self._extractors.values():
            setter = getattr(extractor, "set_page_manifest_path", None)
            if callable(setter):
                setter(path)


def _ocr_image(
    path: Path,
    *,
    language: str = "eng",
    timeout_seconds: int = 120,
    preprocess_mode: OcrPreprocessMode = "none",
    threshold: int = 180,
) -> str:
    tesseract_path = shutil.which("tesseract")
    if tesseract_path is None:
        raise RuntimeError("OCR requires the 'tesseract' executable on PATH")
    with tempfile.TemporaryDirectory() as tmp_dir:
        prepared_path = _prepare_ocr_image(
            path,
            output_dir=Path(tmp_dir),
            preprocess_mode=preprocess_mode,
            threshold=threshold,
        )
        completed = _run_tesseract(
            prepared_path,
            tesseract_path=tesseract_path,
            language=language,
            timeout_seconds=timeout_seconds,
        )
    text = completed.stdout.strip()
    if not text:
        raise ValueError(f"No OCR text found in image: {path}")
    return text


def _ocr_image_result(
    path: Path,
    *,
    language: str = "eng",
    timeout_seconds: int = 120,
    preprocess_mode: OcrPreprocessMode = "none",
    threshold: int = 180,
) -> OcrTextResult:
    tesseract_path = shutil.which("tesseract")
    if tesseract_path is None:
        raise RuntimeError("OCR requires the 'tesseract' executable on PATH")
    with tempfile.TemporaryDirectory() as tmp_dir:
        prepared_path = _prepare_ocr_image(
            path,
            output_dir=Path(tmp_dir),
            preprocess_mode=preprocess_mode,
            threshold=threshold,
        )
        completed = _run_tesseract(
            prepared_path,
            tesseract_path=tesseract_path,
            language=language,
            timeout_seconds=timeout_seconds,
        )
        text = completed.stdout.strip()
        if not text:
            raise ValueError(f"No OCR text found in image: {path}")
        try:
            confidence = _ocr_image_confidence(
                prepared_path,
                language=language,
                timeout_seconds=timeout_seconds,
            )
        except Exception:
            confidence = None
    return OcrTextResult(text=text, confidence=confidence)


def _ocr_image_confidence(
    path: Path,
    *,
    language: str = "eng",
    timeout_seconds: int = 120,
) -> float | None:
    tesseract_path = shutil.which("tesseract")
    if tesseract_path is None:
        raise RuntimeError("OCR requires the 'tesseract' executable on PATH")
    completed = _run_tesseract(
        path,
        tesseract_path=tesseract_path,
        language=language,
        timeout_seconds=timeout_seconds,
        output_format="tsv",
    )
    return parse_tesseract_tsv_confidence(completed.stdout)


def _run_tesseract(
    path: Path,
    *,
    tesseract_path: str,
    language: str,
    timeout_seconds: int,
    output_format: str | None = None,
) -> subprocess.CompletedProcess[str]:
    command = [tesseract_path, str(path), "stdout", "-l", language]
    if output_format is not None:
        command.append(output_format)
    return subprocess.run(  # noqa: S603
        command,
        check=True,
        capture_output=True,
        text=True,
        timeout=timeout_seconds,
    )


def _prepare_ocr_image(
    path: Path,
    *,
    output_dir: Path,
    preprocess_mode: OcrPreprocessMode,
    threshold: int,
) -> Path:
    _validate_ocr_preprocess_config(preprocess_mode, threshold=threshold)
    if preprocess_mode == "none":
        return path
    try:
        image_module = importlib.import_module("PIL.Image")
    except ImportError as exc:
        raise RuntimeError("OCR preprocessing requires installing the 'ocr' extra") from exc

    image = image_module.open(path)
    grayscale = image.convert("L")
    if preprocess_mode == "grayscale":
        prepared = grayscale
    else:
        prepared = _threshold_ocr_image(grayscale, threshold=threshold)
        if preprocess_mode == "deskew":
            prepared = _deskew_ocr_image(prepared, threshold=threshold)
    prepared_path = output_dir / f"{path.stem}.ocr-preprocessed.png"
    prepared.save(prepared_path)
    return prepared_path


def _validate_ocr_preprocess_config(preprocess_mode: str, *, threshold: int) -> None:
    if preprocess_mode not in {"none", "grayscale", "threshold", "deskew"}:
        raise ValueError(f"Unsupported OCR preprocessing mode: {preprocess_mode}")
    if threshold < 0 or threshold > 255:
        raise ValueError("OCR threshold must be between 0 and 255")


def _threshold_ocr_image(image: Any, *, threshold: int) -> Any:
    def binarize(pixel: int) -> int:
        return 255 if pixel > threshold else 0

    return image.point(binarize, mode="1").convert("L")


def _deskew_ocr_image(image: Any, *, threshold: int) -> Any:
    angles = [angle / 2 for angle in range(-10, 11)]
    best_angle = max(angles, key=lambda angle: _ocr_deskew_score(image, angle, threshold=threshold))
    if abs(best_angle) < 0.01:
        return image
    image_module = importlib.import_module("PIL.Image")
    resampling = getattr(image_module, "Resampling", None)
    resample = resampling.BICUBIC if resampling is not None else image_module.BICUBIC
    return image.rotate(
        best_angle,
        resample=resample,
        expand=True,
        fillcolor=255,
    )


def _ocr_deskew_score(image: Any, angle: float, *, threshold: int) -> float:
    rotated = image.rotate(angle, expand=True, fillcolor=255).convert("L")
    width, height = rotated.size
    if width == 0 or height == 0:
        return 0.0
    data = rotated.tobytes()
    row_counts = [
        sum(1 for pixel in data[row * width : (row + 1) * width] if pixel <= threshold)
        for row in range(height)
    ]
    if not any(row_counts):
        return 0.0
    mean = sum(row_counts) / len(row_counts)
    return sum((count - mean) ** 2 for count in row_counts) / len(row_counts)


def parse_tesseract_tsv_confidence(tsv: str) -> float | None:
    lines = [line for line in tsv.splitlines() if line.strip()]
    if not lines:
        return None
    header = lines[0].split("\t")
    try:
        confidence_index = header.index("conf")
        text_index = header.index("text")
    except ValueError:
        return None
    confidences: list[float] = []
    for line in lines[1:]:
        fields = line.split("\t")
        if len(fields) <= max(confidence_index, text_index):
            continue
        if not fields[text_index].strip():
            continue
        try:
            confidence = float(fields[confidence_index])
        except ValueError:
            continue
        if confidence >= 0:
            confidences.append(confidence)
    if not confidences:
        return None
    return sum(confidences) / len(confidences)


def _ocr_pdf(
    path: Path,
    *,
    language: str = "eng",
    timeout_seconds: int = 120,
    dpi: int = 200,
    max_pages: int = 100,
    preprocess_mode: OcrPreprocessMode = "none",
    threshold: int = 180,
) -> str:
    if shutil.which("tesseract") is None:
        raise RuntimeError("Scanned PDF OCR requires the 'tesseract' executable on PATH")
    try:
        pdf2image = importlib.import_module("pdf2image")
    except ImportError as exc:
        raise RuntimeError("Scanned PDF OCR requires installing the 'ocr' extra") from exc

    with tempfile.TemporaryDirectory() as tmp_dir:
        image_paths = cast(
            list[Path],
            pdf2image.convert_from_path(
                path,
                dpi=dpi,
                last_page=max_pages,
                output_folder=tmp_dir,
                fmt="png",
                paths_only=True,
                timeout=timeout_seconds,
            ),
        )
        parts = [
            _ocr_image(
                image_path,
                language=language,
                timeout_seconds=timeout_seconds,
                preprocess_mode=preprocess_mode,
                threshold=threshold,
            )
            for image_path in image_paths
        ]

    text = "\n\n".join(part for part in parts if part.strip())
    if not text:
        raise ValueError(f"No OCR text found in PDF: {path}")
    return text


def _ocr_pdf_page(
    path: Path,
    *,
    page_number: int,
    language: str = "eng",
    timeout_seconds: int = 120,
    dpi: int = 200,
    preprocess_mode: OcrPreprocessMode = "none",
    threshold: int = 180,
    image_artifact_path: Path | None = None,
) -> OcrTextResult:
    if shutil.which("tesseract") is None:
        raise RuntimeError("Scanned PDF OCR requires the 'tesseract' executable on PATH")
    try:
        pdf2image = importlib.import_module("pdf2image")
    except ImportError as exc:
        raise RuntimeError("Scanned PDF OCR requires installing the 'ocr' extra") from exc

    with tempfile.TemporaryDirectory() as tmp_dir:
        image_paths = cast(
            list[Path],
            pdf2image.convert_from_path(
                path,
                dpi=dpi,
                first_page=page_number,
                last_page=page_number,
                output_folder=tmp_dir,
                fmt="png",
                paths_only=True,
                timeout=timeout_seconds,
            ),
        )
        results = [
            _ocr_image_result(
                image_path,
                language=language,
                timeout_seconds=timeout_seconds,
                preprocess_mode=preprocess_mode,
                threshold=threshold,
            )
            for image_path in image_paths
        ]
        preserved_image_path = None
        if image_artifact_path is not None and image_paths:
            _copy_ocr_page_image(image_paths[0], image_artifact_path)
            preserved_image_path = str(image_artifact_path)
    text = "\n\n".join(result.text for result in results if result.text.strip())
    if not text:
        raise ValueError(f"No OCR text found on PDF page {page_number}: {path}")
    confidences = [result.confidence for result in results if result.confidence is not None]
    confidence = sum(confidences) / len(confidences) if confidences else None
    return OcrTextResult(text=text, confidence=confidence, image_path=preserved_image_path)


def _copy_ocr_page_image(source_path: Path, destination_path: Path) -> None:
    _reject_symlinked_output_path(destination_path)
    destination_path.parent.mkdir(parents=True, exist_ok=True)
    temporary_path = destination_path.with_name(f".{destination_path.name}.{uuid.uuid4().hex}.tmp")
    try:
        shutil.copyfile(source_path, temporary_path)
        temporary_path.replace(destination_path)
    except Exception:
        temporary_path.unlink(missing_ok=True)
        raise


def _coerce_ocr_result(value: OcrTextResult | str) -> OcrTextResult:
    if isinstance(value, OcrTextResult):
        return value
    return OcrTextResult(text=value)


def _file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


async def _load_pdf_page_manifest(path: Path | None) -> dict[str, object] | None:
    if path is None or not await asyncio.to_thread(path.exists):
        return None
    if await asyncio.to_thread(path.is_symlink):
        return None
    try:
        manifest_text = await asyncio.to_thread(
            _read_limited_text_file,
            path,
            max_bytes=_MAX_PDF_PAGE_MANIFEST_BYTES,
            label="PDF page manifest",
        )
        payload = json.loads(manifest_text)
    except (OSError, json.JSONDecodeError):
        return None
    return cast(dict[str, object], payload) if isinstance(payload, dict) else None


async def _write_pdf_page_manifest(
    path: Path | None,
    *,
    source_path: Path,
    source_sha256: str,
    extraction_config: dict[str, object],
    page_count: int,
    pages: Sequence[PdfPageExtraction | None],
) -> None:
    if path is None:
        return
    if await asyncio.to_thread(path.is_symlink):
        raise ValueError(f"PDF page manifest path must not be a symlink: {path}")
    records: list[dict[str, object]] = []
    for page_number, page in enumerate(pages, start=1):
        if page is None:
            records.append(
                {
                    "page_number": page_number,
                    "source": "pending",
                    "status": "pending",
                    "attempts": 0,
                    "duration_ms": None,
                    "text": "",
                }
            )
            continue
        records.append(
            {
                "page_number": page.page_number,
                "source": page.source,
                "status": _pdf_page_status(page),
                "chars": len(page.text),
                "raw_chars": len(page.raw_text) if page.raw_text is not None else None,
                "confidence": page.confidence,
                "corrected": page.corrected,
                "error": page.error,
                "warnings": list(page.warnings),
                "attempts": page.attempts,
                "duration_ms": page.duration_ms,
                "raw_text": page.raw_text,
                "corrected_text": page.text if page.corrected else None,
                "image_path": page.image_path,
                "text": page.text,
            }
        )
    payload = json.dumps(
        {
            "generated_by": "librarian",
            "artifact_type": "pdf-page-extraction-manifest",
            "source_path": str(source_path),
            "source_sha256": source_sha256,
            "extraction_config": extraction_config,
            "page_count": page_count,
            "pages": records,
        },
        indent=2,
    )
    await asyncio.to_thread(_write_text_atomic, path, payload)


def _pdf_page_status(page: PdfPageExtraction) -> str:
    if page.error:
        return "failed"
    if page.source == "pending":
        return "pending"
    return "succeeded"


def _write_text_atomic(path: Path, payload: str) -> None:
    _reject_symlinked_output_path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary_path = path.with_name(f".{path.name}.{uuid.uuid4().hex}.tmp")
    try:
        temporary_path.write_text(payload, encoding="utf-8")
        temporary_path.replace(path)
    except Exception:
        temporary_path.unlink(missing_ok=True)
        raise


def _reject_symlinked_output_path(path: Path) -> None:
    if path.is_symlink():
        raise ValueError(f"Output path must not be a symlink: {path}")
    for parent in reversed(path.parents):
        if parent.exists() and parent.is_symlink():
            raise ValueError(f"Output path crosses symlinked parent: {path}")


def _reusable_manifest_pages(
    manifest: dict[str, object] | None,
    *,
    source_sha256: str,
    extraction_config: dict[str, object],
    page_count: int,
) -> dict[int, PdfPageExtraction]:
    if manifest is None:
        return {}
    if manifest.get("source_sha256") != source_sha256:
        return {}
    if manifest.get("extraction_config") != extraction_config:
        return {}
    if manifest.get("page_count") != page_count:
        return {}
    pages_obj = manifest.get("pages")
    if not isinstance(pages_obj, list):
        return {}
    pages: dict[int, PdfPageExtraction] = {}
    for page_obj in cast(list[object], pages_obj):
        if not isinstance(page_obj, dict):
            continue
        item = cast(dict[str, object], page_obj)
        if item.get("status") != "succeeded":
            continue
        page_number = item.get("page_number")
        text = item.get("text")
        source = item.get("source")
        if not isinstance(page_number, int) or not isinstance(text, str):
            continue
        if source not in {"embedded", "ocr"}:
            continue
        pages[page_number] = PdfPageExtraction(
            page_number=page_number,
            text=text,
            source=str(source),
            confidence=cast(float | None, item.get("confidence")),
            corrected=item.get("corrected") is True,
            raw_text=cast(str | None, item.get("raw_text")),
            image_path=cast(str | None, item.get("image_path")),
            warnings=_manifest_warnings(item.get("warnings")),
            attempts=_manifest_attempt_count(item.get("attempts")),
            duration_ms=_manifest_duration_ms(item.get("duration_ms")),
        )
    return pages


def _manifest_page_attempts(manifest: dict[str, object] | None) -> dict[int, int]:
    if manifest is None:
        return {}
    pages_obj = manifest.get("pages")
    if not isinstance(pages_obj, list):
        return {}
    attempts: dict[int, int] = {}
    for page_obj in cast(list[object], pages_obj):
        if not isinstance(page_obj, dict):
            continue
        item = cast(dict[str, object], page_obj)
        page_number = item.get("page_number")
        if isinstance(page_number, int):
            attempts[page_number] = _manifest_attempt_count(item.get("attempts"))
    return attempts


def _manifest_attempt_count(value: object) -> int:
    if isinstance(value, int) and value >= 0:
        return value
    return 0


def _manifest_duration_ms(value: object) -> float | None:
    if isinstance(value, int | float) and value >= 0:
        return float(value)
    return None


def _ocr_page_warnings(
    *,
    confidence: float | None,
    low_confidence_threshold: float,
) -> tuple[str, ...]:
    if confidence is None:
        return ("missing-ocr-confidence",)
    if confidence < low_confidence_threshold:
        return ("low-ocr-confidence",)
    return ()


def _manifest_warnings(value: object) -> tuple[str, ...]:
    if not isinstance(value, list):
        return ()
    warnings: list[str] = []
    for item in cast(list[object], value):
        if isinstance(item, str):
            warnings.append(item)
    return tuple(warnings)


def render_pdf_pages_markdown(path: Path, pages: Sequence[PdfPageExtraction]) -> str:
    """Render ordered PDF page records as canonical Markdown."""
    title = path.stem.replace("_", " ").replace("-", " ").strip() or path.name
    lines = [
        "---",
        "generated_by: librarian",
        "artifact_type: pdf-page-extraction",
        f"source_file: {path.name}",
        f"page_count: {len(pages)}",
        "---",
        "",
        f"# {title}",
        "",
    ]
    for page in sorted(pages, key=lambda item: item.page_number):
        if not page.text.strip() and page.error:
            lines.extend(
                [
                    (
                        f"<!-- page: {page.page_number} source: {page.source} "
                        f"status: failed error: {page.error} -->"
                    ),
                    "",
                ]
            )
            continue
        lines.extend(
            [
                (
                    f"<!-- page: {page.page_number} source: {page.source} "
                    f"corrected: {str(page.corrected).lower()} -->"
                ),
                "",
                f"## Page {page.page_number}",
                "",
                page.text.strip(),
                "",
            ]
        )
    return "\n".join(lines).strip() + "\n"


def _markitdown_worker(
    path: str,
    max_input_bytes: int,
    result_queue: multiprocessing.Queue[tuple[str, str]],
) -> None:
    try:
        resolved_path = Path(path)
        _validate_markitdown_input_size(resolved_path, max_input_bytes)
        result_queue.put(("ok", _extract_markitdown_sync(resolved_path)))
    except Exception as exc:
        result_queue.put(("error", sanitize_error_message(exc)))


def _extract_markitdown_sync(path: Path) -> str:
    try:
        markitdown_module = importlib.import_module("markitdown")
    except ImportError as exc:
        raise RuntimeError(
            "Broad format conversion requires installing the 'universal' extra"
        ) from exc

    markitdown_class = markitdown_module.MarkItDown
    converter = markitdown_class()
    result = converter.convert(str(path))
    text_content = getattr(result, "text_content", None)
    if not isinstance(text_content, str) or not text_content.strip():
        raise ValueError(f"No extractable content found: {path}")
    return text_content


def _validate_markitdown_input_size(path: Path, max_input_bytes: int) -> None:
    _validate_input_size(path, max_input_bytes, "Broad format conversion input")


def _validate_input_size(path: Path, max_input_bytes: int, label: str) -> None:
    byte_size = path.stat().st_size
    if byte_size > max_input_bytes:
        raise ValueError(f"{label} exceeds {max_input_bytes} bytes: {path}")


def _read_limited_text_file(path: Path, *, max_bytes: int, label: str) -> str:
    with path.open("rb") as handle:
        payload = handle.read(max_bytes + 1)
    if len(payload) > max_bytes:
        raise ValueError(f"{label} exceeds {max_bytes} bytes: {path}")
    return payload.decode("utf-8")


def _validate_text_like(path: Path) -> None:
    with path.open("rb") as handle:
        sample = handle.read(_ARCHIVE_SIGNATURE_SAMPLE_BYTES)
    _reject_disallowed_archive_signature(path, sample)
    if b"\x00" in sample:
        raise ValueError(f"Text extraction input appears to be binary: {path}")


def reject_disallowed_archive_signature(path: Path) -> None:
    """Reject archive/container bytes for paths that are not supported container docs."""
    with path.open("rb") as handle:
        sample = handle.read(_ARCHIVE_SIGNATURE_SAMPLE_BYTES)
    _reject_disallowed_archive_signature(path, sample)


def archive_signature_label(payload: bytes) -> str | None:
    for signature, label in ARCHIVE_SIGNATURES:
        if payload.startswith(signature):
            return label
    tar_signature_end = TAR_USTAR_OFFSET + len(b"ustar")
    if (
        len(payload) >= tar_signature_end
        and payload[TAR_USTAR_OFFSET:tar_signature_end] == b"ustar"
    ):
        return "tar"
    return None


def _reject_disallowed_archive_signature(path: Path, payload: bytes) -> None:
    label = archive_signature_label(payload)
    if label is None:
        return
    if label == "zip" and path.suffix.lower() in ZIP_CONTAINER_EXTENSIONS:
        return
    raise ValueError(f"Archive inputs are not supported by default: {label} signature detected")


def _paragraph_text(paragraph: Any) -> str:
    text = str(getattr(paragraph, "text", "")).strip()
    if not text:
        return ""
    style_name = str(getattr(getattr(paragraph, "style", None), "name", "")).lower()
    if "list bullet" in style_name:
        return f"- {text}"
    if "list number" in style_name:
        return f"1. {text}"
    return text


def _table_text(table: Any) -> str:
    rows: list[str] = []
    for row in table.rows:
        cells = [_paragraph_text(cell) for cell in row.cells]
        rendered = " | ".join(cell for cell in cells if cell)
        if rendered:
            rows.append(rendered)
    return "\n".join(rows)
