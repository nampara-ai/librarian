"""Synthetic corpus generation for reproducible benchmark fixtures."""

from __future__ import annotations

import json
import uuid
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any


@dataclass(frozen=True, slots=True)
class SyntheticCorpusResult:
    """Generated synthetic corpus paths."""

    corpus_dir: Path
    suite_path: Path
    files: tuple[Path, ...]
    total_bytes: int
    total_chars: int


def generate_synthetic_corpus(
    *,
    corpus_dir: Path,
    suite_path: Path,
    documents: int = 3,
    paragraphs_per_document: int = 200,
    sentences_per_paragraph: int = 4,
    include_docx: bool = False,
    include_pdf: bool = False,
    include_scanned_pdf: bool = False,
    overwrite: bool = False,
) -> SyntheticCorpusResult:
    """Generate deterministic long-form text fixtures and a corpus-eval suite."""
    if documents < 1:
        raise ValueError("documents must be at least 1")
    if paragraphs_per_document < 1:
        raise ValueError("paragraphs_per_document must be at least 1")
    if sentences_per_paragraph < 1:
        raise ValueError("sentences_per_paragraph must be at least 1")
    _reject_symlinked_path(corpus_dir, label="Corpus directory")
    _reject_symlinked_path(suite_path, label="Corpus eval suite path")
    corpus_dir.mkdir(parents=True, exist_ok=True)
    suite_path.parent.mkdir(parents=True, exist_ok=True)
    if suite_path.exists() and not overwrite:
        raise FileExistsError(f"Corpus eval suite already exists: {suite_path}")
    if suite_path.is_symlink():
        raise ValueError(f"Corpus eval suite path must not be a symlink: {suite_path}")
    files: list[Path] = []
    cases: list[dict[str, object]] = []
    total_bytes = 0
    total_chars = 0
    topics = (
        (
            "equine-training",
            "Horse Training Seminar",
            "canter transitions",
            "saddle fit",
            "636",
        ),
        (
            "health-interview",
            "Health Interview Transcript",
            "medicine adherence",
            "follow-up care",
            "610",
        ),
        (
            "library-workshop",
            "Library Workshop Transcript",
            "catalog metadata",
            "search recall",
            "020",
        ),
    )
    for index in range(documents):
        slug, title, first_phrase, second_phrase, classification_prefix = topics[
            index % len(topics)
        ]
        path = corpus_dir / f"{index + 1:03d}-{slug}.md"
        if path.exists() and not overwrite:
            raise FileExistsError(f"Corpus file already exists: {path}")
        if path.is_symlink():
            raise ValueError(f"Corpus file path must not be a symlink: {path}")
        text = _synthetic_markdown(
            title=title,
            first_phrase=first_phrase,
            second_phrase=second_phrase,
            paragraphs=paragraphs_per_document,
            sentences_per_paragraph=sentences_per_paragraph,
            document_number=index + 1,
        )
        _write_text_atomic(path, text)
        files.append(path)
        total_bytes += path.stat().st_size
        total_chars += len(text)
        cases.append(
            {
                "name": f"{title} {index + 1}",
                "source_path": str(path.relative_to(suite_path.parent)),
                "tags": ["synthetic", "long-document", slug],
                "format": "md",
                "process": True,
                "expected_contains": [first_phrase, second_phrase],
                "expected_search_phrases": [first_phrase, second_phrase],
                "expected_classification_prefix": classification_prefix,
                "min_output_char_ratio": 0.5,
                "max_output_char_ratio": 4.0,
                "require_markdown_headings": True,
            }
        )
    if include_docx:
        for index, (slug, title, first_phrase, second_phrase, classification_prefix) in enumerate(
            topics,
            start=1,
        ):
            path = corpus_dir / f"{documents + index:03d}-{slug}-docx.docx"
            if path.exists() and not overwrite:
                raise FileExistsError(f"Corpus file already exists: {path}")
            if path.is_symlink():
                raise ValueError(f"Corpus file path must not be a symlink: {path}")
            doc_text = _write_synthetic_docx_atomic(
                path,
                title=title,
                first_phrase=first_phrase,
                second_phrase=second_phrase,
                paragraphs=max(3, min(paragraphs_per_document, 25)),
                document_number=documents + index,
            )
            files.append(path)
            total_bytes += path.stat().st_size
            total_chars += len(doc_text)
            cases.append(
                {
                    "name": f"{title} DOCX {index}",
                    "source_path": str(path.relative_to(suite_path.parent)),
                    "tags": ["synthetic", "docx", "tables", "headers-footers", slug],
                    "format": "md",
                    "process": True,
                    "expected_contains": [
                        first_phrase,
                        second_phrase,
                        "Synthetic header",
                        "Synthetic footer",
                        "Table checkpoint",
                    ],
                    "expected_search_phrases": [first_phrase, "Table checkpoint"],
                    "expected_classification_prefix": classification_prefix,
                    "min_output_char_ratio": 0.02,
                    "max_output_char_ratio": 20.0,
                }
            )
    if include_pdf:
        base_index = len(files)
        for index, (slug, title, first_phrase, second_phrase, classification_prefix) in enumerate(
            topics,
            start=1,
        ):
            path = corpus_dir / f"{base_index + index:03d}-{slug}-embedded.pdf"
            if path.exists() and not overwrite:
                raise FileExistsError(f"Corpus file already exists: {path}")
            if path.is_symlink():
                raise ValueError(f"Corpus file path must not be a symlink: {path}")
            pdf_text, page_count = _write_synthetic_pdf_atomic(
                path,
                title=title,
                first_phrase=first_phrase,
                second_phrase=second_phrase,
                paragraphs=max(2, min(paragraphs_per_document, 50)),
                document_number=base_index + index,
            )
            files.append(path)
            total_bytes += path.stat().st_size
            total_chars += len(pdf_text)
            cases.append(
                {
                    "name": f"{title} Embedded PDF {index}",
                    "source_path": str(path.relative_to(suite_path.parent)),
                    "tags": ["synthetic", "pdf", "embedded-text", slug],
                    "format": "md",
                    "process": True,
                    "expected_contains": [first_phrase, second_phrase],
                    "expected_search_phrases": [first_phrase, second_phrase],
                    "expected_classification_prefix": classification_prefix,
                    "expected_page_count": page_count,
                    "min_output_char_ratio": 0.05,
                    "max_output_char_ratio": 50.0,
                    "require_markdown_headings": True,
                }
            )
    if include_scanned_pdf:
        base_index = len(files)
        for index, (slug, title, first_phrase, second_phrase, classification_prefix) in enumerate(
            topics[:2],
            start=1,
        ):
            path = corpus_dir / f"{base_index + index:03d}-{slug}-scanned.pdf"
            if path.exists() and not overwrite:
                raise FileExistsError(f"Corpus file already exists: {path}")
            if path.is_symlink():
                raise ValueError(f"Corpus file path must not be a symlink: {path}")
            pdf_text, page_count = _write_synthetic_scanned_pdf_atomic(
                path,
                title=title,
                first_phrase=first_phrase,
                second_phrase=second_phrase,
                paragraphs=max(2, min(paragraphs_per_document, 10)),
                document_number=base_index + index,
                mixed=index == 2,
            )
            files.append(path)
            total_bytes += path.stat().st_size
            total_chars += len(pdf_text)
            tags = ["synthetic", "pdf", "scanned", "ocr", slug]
            if index == 2:
                tags.append("mixed-embedded-scanned")
            cases.append(
                {
                    "name": f"{title} OCR PDF {index}",
                    "source_path": str(path.relative_to(suite_path.parent)),
                    "tags": tags,
                    "format": "md",
                    "process": True,
                    "expected_contains": [first_phrase, second_phrase],
                    "expected_search_phrases": [first_phrase],
                    "expected_classification_prefix": classification_prefix,
                    "expected_page_count": page_count,
                    "min_output_char_ratio": 0.001,
                    "max_output_char_ratio": 50.0,
                    "require_markdown_headings": True,
                }
            )
    _write_text_atomic(suite_path, json.dumps({"cases": cases}, indent=2) + "\n")
    return SyntheticCorpusResult(
        corpus_dir=corpus_dir,
        suite_path=suite_path,
        files=tuple(files),
        total_bytes=total_bytes,
        total_chars=total_chars,
    )


def _write_text_atomic(path: Path, payload: str) -> None:
    _reject_symlinked_path(path, label="Output path")
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary_path = path.with_name(f".{path.name}.{uuid.uuid4().hex}.tmp")
    try:
        temporary_path.write_text(payload, encoding="utf-8")
        temporary_path.replace(path)
    except Exception:
        temporary_path.unlink(missing_ok=True)
        raise


def _write_synthetic_docx_atomic(
    path: Path,
    *,
    title: str,
    first_phrase: str,
    second_phrase: str,
    paragraphs: int,
    document_number: int,
) -> str:
    _reject_symlinked_path(path, label="Output path")
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary_path = path.with_name(f".{path.name}.{uuid.uuid4().hex}.tmp.docx")
    try:
        from docx import Document

        document = Document()
        section = document.sections[0]
        section.header.paragraphs[0].text = (
            f"Synthetic header for {title} covering {first_phrase}"
        )
        document.add_heading(title, level=1)
        expected_parts = [
            section.header.paragraphs[0].text,
            title,
        ]
        for paragraph in range(1, paragraphs + 1):
            text = (
                f"Document {document_number} paragraph {paragraph} discusses "
                f"{first_phrase}, {second_phrase}, source fidelity, and review checkpoints."
            )
            document.add_paragraph(text)
            expected_parts.append(text)
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Table checkpoint"
        table.cell(0, 1).text = first_phrase
        table.cell(1, 0).text = "Review topic"
        table.cell(1, 1).text = second_phrase
        expected_parts.extend(
            (
                "Table checkpoint",
                first_phrase,
                "Review topic",
                second_phrase,
            )
        )
        section.footer.paragraphs[0].text = (
            f"Synthetic footer for {title} preserving {second_phrase}"
        )
        expected_parts.append(section.footer.paragraphs[0].text)
        document.save(str(temporary_path))
        temporary_path.replace(path)
        return "\n".join(expected_parts)
    except Exception:
        temporary_path.unlink(missing_ok=True)
        raise


def _write_synthetic_pdf_atomic(
    path: Path,
    *,
    title: str,
    first_phrase: str,
    second_phrase: str,
    paragraphs: int,
    document_number: int,
) -> tuple[str, int]:
    _reject_symlinked_path(path, label="Output path")
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary_path = path.with_name(f".{path.name}.{uuid.uuid4().hex}.tmp.pdf")
    try:
        text = _synthetic_plain_text(
            title=title,
            first_phrase=first_phrase,
            second_phrase=second_phrase,
            paragraphs=paragraphs,
            document_number=document_number,
        )
        pages = _paginate_pdf_lines(text.splitlines(), lines_per_page=34)
        temporary_path.write_bytes(_render_simple_pdf(pages))
        temporary_path.replace(path)
        return text, len(pages)
    except Exception:
        temporary_path.unlink(missing_ok=True)
        raise


def _write_synthetic_scanned_pdf_atomic(
    path: Path,
    *,
    title: str,
    first_phrase: str,
    second_phrase: str,
    paragraphs: int,
    document_number: int,
    mixed: bool,
) -> tuple[str, int]:
    _reject_symlinked_path(path, label="Output path")
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary_path = path.with_name(f".{path.name}.{uuid.uuid4().hex}.tmp.pdf")
    try:
        text = _synthetic_plain_text(
            title=title,
            first_phrase=first_phrase,
            second_phrase=second_phrase,
            paragraphs=paragraphs,
            document_number=document_number,
        )
        lines = [line for line in text.splitlines() if line.strip()]
        scanned_pages = _paginate_pdf_lines(lines, lines_per_page=8)
        if mixed:
            embedded = [scanned_pages[0]]
            scanned = scanned_pages[1:] or scanned_pages[:1]
            payload = _render_mixed_text_image_pdf(embedded, scanned)
        else:
            payload = _render_image_only_pdf(scanned_pages)
        temporary_path.write_bytes(payload)
        temporary_path.replace(path)
        return text, len(scanned_pages)
    except Exception:
        temporary_path.unlink(missing_ok=True)
        raise


def _reject_symlinked_path(path: Path, *, label: str) -> None:
    for current in (*reversed(path.parents), path):
        if current.exists() and current.is_symlink():
            raise ValueError(f"{label} crosses symlink: {path}")


def _synthetic_markdown(
    *,
    title: str,
    first_phrase: str,
    second_phrase: str,
    paragraphs: int,
    sentences_per_paragraph: int,
    document_number: int,
) -> str:
    lines = [f"# {title}", ""]
    for paragraph in range(1, paragraphs + 1):
        section = ((paragraph - 1) // 25) + 1
        if paragraph == 1 or (paragraph - 1) % 25 == 0:
            lines.extend([f"## Section {section}", ""])
        sentences = [
            (
                f"Speaker {paragraph % 4 + 1}: In document {document_number}, "
                f"paragraph {paragraph}, we discuss {first_phrase}, {second_phrase}, "
                "source fidelity, page markers, and searchable transcript structure."
            )
        ]
        for sentence in range(2, sentences_per_paragraph + 1):
            sentences.append(
                " ".join(
                    (
                        f"The deterministic note {sentence} tracks review topic",
                        f"{(paragraph + sentence) % 17},",
                        "keeps stable wording for benchmark diffs, and preserves",
                        "enough transcript-like repetition to exercise chunking.",
                    )
                )
            )
        lines.extend((" ".join(sentences), ""))
    return "\n".join(lines)


def _synthetic_plain_text(
    *,
    title: str,
    first_phrase: str,
    second_phrase: str,
    paragraphs: int,
    document_number: int,
) -> str:
    lines = [title, ""]
    for paragraph in range(1, paragraphs + 1):
        lines.append(
            f"Document {document_number} paragraph {paragraph} discusses {first_phrase}, "
            f"{second_phrase}, source fidelity, page markers, and searchable transcript structure."
        )
    return "\n".join(lines)


def _paginate_pdf_lines(lines: list[str], *, lines_per_page: int) -> list[list[str]]:
    if not lines:
        return [[]]
    return [
        lines[index : index + lines_per_page]
        for index in range(0, len(lines), lines_per_page)
    ]


def _render_simple_pdf(pages: list[list[str]]) -> bytes:
    objects: list[bytes] = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    page_refs: list[str] = []
    for page_index, page_lines in enumerate(pages):
        page_object = len(objects) + 1
        content_object = page_object + 1
        page_refs.append(f"{page_object} 0 R")
        content = _pdf_page_content(page_lines)
        objects.append(
            (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                f"/Resources << /Font << /F1 3 0 R >> >> /Contents {content_object} 0 R >>"
            ).encode("ascii")
        )
        objects.append(
            f"<< /Length {len(content)} >>\nstream\n".encode("ascii")
            + content
            + b"\nendstream"
        )
        del page_index
    objects[1] = (
        f"<< /Type /Pages /Kids [{' '.join(page_refs)}] /Count {len(page_refs)} >>"
    ).encode("ascii")
    return _assemble_pdf(objects)


def _render_image_only_pdf(pages: list[list[str]]) -> bytes:
    objects: list[bytes] = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"",
    ]
    page_refs: list[str] = []
    for page_lines in pages:
        page_object = len(objects) + 1
        content_object = page_object + 1
        image_object = page_object + 2
        page_refs.append(f"{page_object} 0 R")
        image = _render_scan_page_jpeg(page_lines)
        content = f"q\n468 0 0 648 72 72 cm\n/Im{image_object} Do\nQ".encode("ascii")
        objects.append(
            (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                f"/Resources << /XObject << /Im{image_object} {image_object} 0 R >> >> "
                f"/Contents {content_object} 0 R >>"
            ).encode("ascii")
        )
        objects.append(
            f"<< /Length {len(content)} >>\nstream\n".encode("ascii")
            + content
            + b"\nendstream"
        )
        objects.append(
            (
                f"<< /Type /XObject /Subtype /Image /Width 1240 /Height 1754 "
                f"/ColorSpace /DeviceRGB /BitsPerComponent 8 /Filter /DCTDecode "
                f"/Length {len(image)} >>\nstream\n"
            ).encode("ascii")
            + image
            + b"\nendstream"
        )
    objects[1] = (
        f"<< /Type /Pages /Kids [{' '.join(page_refs)}] /Count {len(page_refs)} >>"
    ).encode("ascii")
    return _assemble_pdf(objects)


def _render_mixed_text_image_pdf(
    embedded_pages: list[list[str]],
    scanned_pages: list[list[str]],
) -> bytes:
    objects: list[bytes] = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    page_refs: list[str] = []
    for page_lines in embedded_pages:
        page_object = len(objects) + 1
        content_object = page_object + 1
        page_refs.append(f"{page_object} 0 R")
        content = _pdf_page_content(page_lines)
        objects.append(
            (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                f"/Resources << /Font << /F1 3 0 R >> >> /Contents {content_object} 0 R >>"
            ).encode("ascii")
        )
        objects.append(
            f"<< /Length {len(content)} >>\nstream\n".encode("ascii")
            + content
            + b"\nendstream"
        )
    image_pdf_objects = _image_page_objects(scanned_pages, first_object=len(objects) + 1)
    page_refs.extend(image_pdf_objects.page_refs)
    objects.extend(image_pdf_objects.objects)
    objects[1] = (
        f"<< /Type /Pages /Kids [{' '.join(page_refs)}] /Count {len(page_refs)} >>"
    ).encode("ascii")
    return _assemble_pdf(objects)


@dataclass(frozen=True, slots=True)
class _ImagePdfObjects:
    objects: list[bytes]
    page_refs: list[str]


def _image_page_objects(pages: list[list[str]], *, first_object: int) -> _ImagePdfObjects:
    objects: list[bytes] = []
    page_refs: list[str] = []
    next_object = first_object
    for page_lines in pages:
        page_object = next_object
        content_object = page_object + 1
        image_object = page_object + 2
        next_object += 3
        page_refs.append(f"{page_object} 0 R")
        image = _render_scan_page_jpeg(page_lines)
        content = f"q\n468 0 0 648 72 72 cm\n/Im{image_object} Do\nQ".encode("ascii")
        objects.append(
            (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                f"/Resources << /XObject << /Im{image_object} {image_object} 0 R >> >> "
                f"/Contents {content_object} 0 R >>"
            ).encode("ascii")
        )
        objects.append(
            f"<< /Length {len(content)} >>\nstream\n".encode("ascii")
            + content
            + b"\nendstream"
        )
        objects.append(
            (
                f"<< /Type /XObject /Subtype /Image /Width 1240 /Height 1754 "
                f"/ColorSpace /DeviceRGB /BitsPerComponent 8 /Filter /DCTDecode "
                f"/Length {len(image)} >>\nstream\n"
            ).encode("ascii")
            + image
            + b"\nendstream"
        )
    return _ImagePdfObjects(objects=objects, page_refs=page_refs)


def _render_scan_page_jpeg(lines: list[str]) -> bytes:
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError as exc:
        raise RuntimeError("Scanned PDF fixtures require Pillow via the 'ocr' extra") from exc

    image = Image.new("RGB", (1240, 1754), "white")
    draw = ImageDraw.Draw(image)
    font: Any = _load_scan_font(ImageFont)
    y = 140
    for line in lines:
        draw.text((120, y), line[:74], fill="black", font=font)
        y += 72
    buffer = BytesIO()
    image.save(buffer, format="JPEG", quality=92)
    return buffer.getvalue()


def _load_scan_font(image_font_module: Any) -> Any:
    font_paths = (
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/Library/Fonts/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Arial.ttf",
    )
    for font_path in font_paths:
        try:
            return image_font_module.truetype(font_path, 38)
        except OSError:
            continue
    return image_font_module.load_default()


def _pdf_page_content(lines: list[str]) -> bytes:
    commands = ["BT", "/F1 10 Tf", "72 740 Td", "14 TL"]
    for index, line in enumerate(lines):
        if index:
            commands.append("T*")
        commands.append(f"({_pdf_escape(line[:92])}) Tj")
    commands.append("ET")
    return "\n".join(commands).encode("ascii")


def _pdf_escape(value: str) -> str:
    safe = value.encode("ascii", errors="ignore").decode("ascii")
    return safe.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _assemble_pdf(objects: list[bytes]) -> bytes:
    output = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for number, payload in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{number} 0 obj\n".encode("ascii"))
        output.extend(payload)
        output.extend(b"\nendobj\n")
    xref_offset = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    return bytes(output)
