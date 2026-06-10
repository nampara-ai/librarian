import json
import multiprocessing
import shutil
import subprocess
from pathlib import Path
from typing import Any, cast

import pytest

from librarian.ingest import extractors
from librarian.ingest.extractors import (
    CompositeExtractor,
    ImageOcrExtractor,
    MarkItDownExtractor,
    OcrTextResult,
    PdfExtractor,
    TextFamilyExtractor,
    TranscriptFileExtractor,
)
from librarian.llm.mock import MockLLMProvider


@pytest.mark.asyncio
async def test_text_family_extractor_reads_markdown(tmp_path: Path) -> None:
    path = tmp_path / "notes.md"
    path.write_text("# Notes\n\nTranscript text", encoding="utf-8")

    text = await TextFamilyExtractor().extract(path)

    assert "Transcript text" in text


@pytest.mark.asyncio
async def test_text_family_extractor_uses_bounded_read(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    path = tmp_path / "notes.md"
    path.write_text("# Notes\n\nTranscript text", encoding="utf-8")

    def fail_read_text(self: Path, *args: Any, **kwargs: Any) -> str:
        del args, kwargs
        raise AssertionError(f"unbounded read_text called for {self}")

    monkeypatch.setattr(Path, "read_text", fail_read_text)

    text = await TextFamilyExtractor(max_input_bytes=100).extract(path)

    assert "Transcript text" in text


@pytest.mark.asyncio
async def test_text_family_extractor_rejects_large_inputs(tmp_path: Path) -> None:
    path = tmp_path / "notes.txt"
    path.write_text("too large", encoding="utf-8")

    with pytest.raises(ValueError, match="Text extraction input exceeds"):
        await TextFamilyExtractor(max_input_bytes=4).extract(path)


@pytest.mark.asyncio
async def test_text_family_extractor_pretty_prints_json(tmp_path: Path) -> None:
    path = tmp_path / "notes.json"
    path.write_text(json.dumps({"speaker": "A", "text": "hello"}), encoding="utf-8")

    text = await TextFamilyExtractor().extract(path)

    assert '"speaker": "A"' in text
    assert "\n" in text


@pytest.mark.asyncio
async def test_text_family_extractor_returns_invalid_json_as_text(tmp_path: Path) -> None:
    path = tmp_path / "broken.json"
    path.write_text("{not json", encoding="utf-8")

    text = await TextFamilyExtractor().extract(path)

    assert text == "{not json"


@pytest.mark.asyncio
async def test_text_family_extractor_rejects_binary_content(tmp_path: Path) -> None:
    path = tmp_path / "binary.txt"
    path.write_bytes(b"text\x00binary")

    with pytest.raises(ValueError, match="appears to be binary"):
        await TextFamilyExtractor().extract(path)


@pytest.mark.asyncio
async def test_text_family_extractor_rejects_renamed_zip_archive(tmp_path: Path) -> None:
    path = tmp_path / "renamed.txt"
    path.write_bytes(b"PK\x03\x04renamed zip archive")

    with pytest.raises(ValueError, match="Archive inputs are not supported"):
        await TextFamilyExtractor().extract(path)


@pytest.mark.asyncio
async def test_transcript_file_extractor_normalizes_srt(tmp_path: Path) -> None:
    path = tmp_path / "captions.srt"
    path.write_text(
        "1\n"
        "00:00:01,000 --> 00:00:02,000\n"
        "Ada: Hello\n\n"
        "2\n"
        "00:00:02,000 --> 00:00:04,000\n"
        "world.\n",
        encoding="utf-8",
    )

    text = await TranscriptFileExtractor().extract(path)

    assert text == "- [00:01] Ada: Hello world."


@pytest.mark.asyncio
async def test_composite_extractor_normalizes_vtt(tmp_path: Path) -> None:
    path = tmp_path / "captions.vtt"
    path.write_text(
        "WEBVTT\n\n"
        "00:00:00.000 --> 00:00:01.000 align:start\n"
        "Opening line\n\n"
        "00:00:01.000 --> 00:00:03.000\n"
        "continues.\n",
        encoding="utf-8",
    )

    text = await CompositeExtractor().extract(path)

    assert text == "- [00:00] Opening line continues."


@pytest.mark.asyncio
async def test_transcript_file_extractor_rejects_malformed_captions(tmp_path: Path) -> None:
    path = tmp_path / "captions.srt"
    path.write_text("no timestamps here", encoding="utf-8")

    with pytest.raises(ValueError, match="No timestamped transcript segments found"):
        await TranscriptFileExtractor().extract(path)


@pytest.mark.asyncio
async def test_composite_extractor_rejects_unknown_extension(tmp_path: Path) -> None:
    path = tmp_path / "notes.bin"
    path.write_bytes(b"binary")

    with pytest.raises(ValueError, match="Unsupported file extension"):
        await CompositeExtractor().extract(path)


@pytest.mark.asyncio
async def test_composite_extractor_rejects_zip_archives(tmp_path: Path) -> None:
    path = tmp_path / "archive.zip"
    path.write_bytes(b"PK")

    with pytest.raises(ValueError, match="Archive inputs are not supported"):
        await CompositeExtractor().extract(path)


@pytest.mark.asyncio
async def test_markitdown_extractor_rejects_large_inputs_before_import(tmp_path: Path) -> None:
    path = tmp_path / "fixture.html"
    path.write_text("<p>too large</p>", encoding="utf-8")

    with pytest.raises(ValueError, match="exceeds"):
        await MarkItDownExtractor(max_input_bytes=4).extract(path)


@pytest.mark.asyncio
async def test_markitdown_extractor_rejects_renamed_zip_archive_before_import(
    tmp_path: Path,
) -> None:
    path = tmp_path / "fixture.html"
    path.write_bytes(b"PK\x03\x04renamed zip archive")

    with pytest.raises(ValueError, match="Archive inputs are not supported"):
        await MarkItDownExtractor().extract(path)


@pytest.mark.asyncio
async def test_markitdown_extractor_times_out(tmp_path: Path) -> None:
    path = tmp_path / "fixture.html"
    path.write_text("<p>small but bounded</p>", encoding="utf-8")

    with pytest.raises(TimeoutError, match="timed out"):
        await MarkItDownExtractor(timeout_seconds=0).extract(path)


def test_markitdown_worker_redacts_child_process_errors(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    path = tmp_path / "fixture.html"
    path.write_text("<p>small but failing</p>", encoding="utf-8")

    def fail_extract(path: Path) -> str:
        del path
        raise RuntimeError("markitdown failed api_key=abc123 sk-testSECRET123")

    monkeypatch.setattr(extractors, "_extract_markitdown_sync", fail_extract)
    queue: multiprocessing.Queue[tuple[str, str]] = multiprocessing.Queue(maxsize=1)

    cast(Any, extractors)._markitdown_worker(str(path), 1024, queue)

    status, payload = queue.get(timeout=1)
    assert status == "error"
    assert payload == "markitdown failed api_key=[REDACTED] [REDACTED]"
    assert "abc123" not in payload
    assert "sk-testSECRET123" not in payload


@pytest.mark.asyncio
async def test_docx_extractor_reads_fixture(tmp_path: Path) -> None:
    from docx import Document

    path = tmp_path / "fixture.docx"
    document = Document()
    document.add_paragraph("DOCX fixture text")
    document.save(str(path))

    text = await CompositeExtractor().extract(path)

    assert "DOCX fixture text" in text


@pytest.mark.asyncio
async def test_docx_extractor_reads_tables_headers_and_footers(tmp_path: Path) -> None:
    from docx import Document

    path = tmp_path / "fixture.docx"
    document = Document()
    document.sections[0].header.paragraphs[0].text = "Header text"
    document.add_paragraph("Body paragraph")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Table left"
    table.cell(0, 1).text = "Table right"
    document.sections[0].footer.paragraphs[0].text = "Footer text"
    document.save(str(path))

    text = await CompositeExtractor().extract(path)

    assert "Header text" in text
    assert "Body paragraph" in text
    assert "Table left | Table right" in text
    assert "Footer text" in text


@pytest.mark.asyncio
async def test_docx_extractor_preserves_list_markers(tmp_path: Path) -> None:
    from docx import Document

    path = tmp_path / "fixture.docx"
    document = Document()
    document.add_paragraph("Bullet item", style="List Bullet")
    document.add_paragraph("Numbered item", style="List Number")
    document.save(str(path))

    text = await CompositeExtractor().extract(path)

    assert "- Bullet item" in text
    assert "1. Numbered item" in text


@pytest.mark.skipif(shutil.which("tesseract") is None, reason="tesseract not installed")
@pytest.mark.asyncio
async def test_image_ocr_extractor_reads_fixture(tmp_path: Path) -> None:
    pytest.importorskip("PIL", reason="pillow not installed")
    from PIL import Image, ImageDraw, ImageFont

    path = tmp_path / "fixture.png"
    image = Image.new("RGB", (800, 220), "white")
    draw = ImageDraw.Draw(image)
    font = ImageFont.load_default(size=48)
    draw.text((40, 70), "OCR fixture text", fill="black", font=font)
    image.save(path)

    text = await CompositeExtractor().extract(path)

    assert "fixture text" in text


@pytest.mark.asyncio
async def test_image_ocr_extractor_passes_timeout(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    path = tmp_path / "fixture.png"
    path.write_bytes(b"not real image")
    captured_kwargs: dict[str, Any] = {}

    def fake_run(cmd: list[str], **kwargs: Any) -> subprocess.CompletedProcess[str]:
        captured_kwargs.update(kwargs)
        return subprocess.CompletedProcess(args=cmd, returncode=0, stdout="OCR text\n")

    def fake_which(name: str) -> str:
        return "/usr/bin/tesseract"

    monkeypatch.setattr(shutil, "which", fake_which)
    monkeypatch.setattr(extractors.subprocess, "run", fake_run)

    text = await ImageOcrExtractor(timeout_seconds=7).extract(path)

    assert text == "OCR text"
    assert captured_kwargs["timeout"] == 7


@pytest.mark.asyncio
async def test_pdf_ocr_passes_timeout_to_rasterizer(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    path = tmp_path / "fixture.pdf"
    path.write_bytes(b"%PDF")
    captured_kwargs: dict[str, Any] = {}

    class FakePdf2Image:
        @staticmethod
        def convert_from_path(*args: Any, **kwargs: Any) -> list[Path]:
            del args
            captured_kwargs.update(kwargs)
            return [tmp_path / "page.png"]

    class FakePage:
        @staticmethod
        def extract_text() -> None:
            return None

    class FakePdf:
        pages = [FakePage()]

        def __enter__(self) -> "FakePdf":
            return self

        def __exit__(self, *args: object) -> None:
            del args

    class FakePdfPlumber:
        @staticmethod
        def open(path: Path) -> FakePdf:
            del path
            return FakePdf()

    def fake_import_module(name: str) -> object:
        if name == "pdfplumber":
            return FakePdfPlumber
        if name == "pdf2image":
            return FakePdf2Image
        return __import__(name)

    def fake_which(name: str) -> str:
        return f"/usr/bin/{name}"

    def fake_ocr_image_result(*args: object, **kwargs: object) -> OcrTextResult:
        del args, kwargs
        return OcrTextResult(text="OCR text")

    monkeypatch.setattr(shutil, "which", fake_which)
    monkeypatch.setattr(extractors.importlib, "import_module", fake_import_module)
    monkeypatch.setattr("librarian.ingest.extractors._ocr_image_result", fake_ocr_image_result)

    text = await PdfExtractor(ocr_timeout_seconds=9, ocr_correction_mode="never").extract(path)

    assert "## Page 1" in text
    assert "OCR text" in text
    assert "source: ocr" in text
    assert captured_kwargs["timeout"] == 9


@pytest.mark.asyncio
async def test_pdf_extractor_ocr_handles_mixed_text_and_scanned_pages(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    path = tmp_path / "fixture.pdf"
    path.write_bytes(b"%PDF")

    class FakePage:
        def __init__(self, text: str | None) -> None:
            self._text = text

        def extract_text(self) -> str | None:
            return self._text

    class FakePdf:
        pages = [FakePage("Text page 1"), FakePage(None), FakePage("Text page 3")]

        def __enter__(self) -> "FakePdf":
            return self

        def __exit__(self, *args: object) -> None:
            del args

    class FakePdfPlumber:
        @staticmethod
        def open(path: Path) -> FakePdf:
            del path
            return FakePdf()

    def fake_import_module(name: str) -> object:
        if name == "pdfplumber":
            return FakePdfPlumber
        return __import__(name)

    def fake_ocr_pdf_page(*args: object, **kwargs: object) -> str:
        del args, kwargs
        return "OCR page 2"

    monkeypatch.setattr(extractors.importlib, "import_module", fake_import_module)
    monkeypatch.setattr("librarian.ingest.extractors._ocr_pdf_page", fake_ocr_pdf_page)

    text = await PdfExtractor(ocr_correction_mode="never").extract(path)

    assert "## Page 1" in text
    assert "Text page 1" in text
    assert "## Page 2" in text
    assert "OCR page 2" in text
    assert "## Page 3" in text
    assert "Text page 3" in text


@pytest.mark.asyncio
async def test_pdf_extractor_writes_and_reuses_page_manifest(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    path = tmp_path / "fixture.pdf"
    manifest = tmp_path / "fixture.md.pages.json"
    path.write_bytes(b"%PDF stable")

    class FakePage:
        def __init__(self, text: str | None) -> None:
            self._text = text

        def extract_text(self) -> str | None:
            return self._text

    class FakePdf:
        pages = [FakePage("Text page 1"), FakePage(None)]

        def __enter__(self) -> "FakePdf":
            return self

        def __exit__(self, *args: object) -> None:
            del args

    class FakePdfPlumber:
        @staticmethod
        def open(path: Path) -> FakePdf:
            del path
            return FakePdf()

    def fake_import_module(name: str) -> object:
        if name == "pdfplumber":
            return FakePdfPlumber
        return __import__(name)

    calls = 0

    def fake_ocr_pdf_page(*args: object, **kwargs: object) -> str:
        del args, kwargs
        nonlocal calls
        calls += 1
        return "OCR page 2"

    monkeypatch.setattr(extractors.importlib, "import_module", fake_import_module)
    monkeypatch.setattr("librarian.ingest.extractors._ocr_pdf_page", fake_ocr_pdf_page)
    first = PdfExtractor(ocr_correction_mode="never")
    first.set_page_manifest_path(manifest)

    text = await first.extract(path)

    payload = json.loads(manifest.read_text(encoding="utf-8"))
    assert calls == 1
    assert payload["artifact_type"] == "pdf-page-extraction-manifest"
    assert payload["schema_version"] == 1
    assert payload["summary"] == {
        "status": "succeeded",
        "status_counts": {"failed": 0, "pending": 0, "succeeded": 2},
        "source_counts": {"embedded": 1, "ocr": 1},
        "warning_counts": {"missing-ocr-confidence": 1},
        "attempts": 1,
        "ocr_pages": 1,
        "corrected_pages": 0,
        "average_ocr_confidence": None,
        "max_page_duration_ms": payload["summary"]["max_page_duration_ms"],
    }
    assert isinstance(payload["summary"]["max_page_duration_ms"], float)
    assert payload["pages"][1]["text"] == "OCR page 2"
    assert payload["pages"][1]["raw_text"] == "OCR page 2"
    assert payload["pages"][1]["corrected_text"] is None
    assert "OCR page 2" in text

    def fail_ocr_pdf_page(*args: object, **kwargs: object) -> str:
        del args, kwargs
        raise AssertionError("manifest should avoid replaying OCR")

    monkeypatch.setattr("librarian.ingest.extractors._ocr_pdf_page", fail_ocr_pdf_page)
    second = PdfExtractor(ocr_correction_mode="never")
    second.set_page_manifest_path(manifest)

    resumed = await second.extract(path)

    assert "OCR page 2" in resumed


@pytest.mark.asyncio
async def test_pdf_extractor_page_manifest_tracks_pending_failed_and_retry_attempts(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    path = tmp_path / "fixture.pdf"
    manifest = tmp_path / "fixture.md.pages.json"
    path.write_bytes(b"%PDF retry")

    class FakePage:
        def __init__(self, text: str | None) -> None:
            self._text = text

        def extract_text(self) -> str | None:
            return self._text

    class FakePdf:
        pages = [FakePage("Text page 1"), FakePage(None)]

        def __enter__(self) -> "FakePdf":
            return self

        def __exit__(self, *args: object) -> None:
            del args

    class FakePdfPlumber:
        @staticmethod
        def open(path: Path) -> FakePdf:
            del path
            return FakePdf()

    def fake_import_module(name: str) -> object:
        if name == "pdfplumber":
            return FakePdfPlumber
        return __import__(name)

    calls = 0

    def flaky_ocr_pdf_page(*args: object, **kwargs: object) -> str:
        del args, kwargs
        nonlocal calls
        calls += 1
        if calls == 1:
            pending_payload = json.loads(manifest.read_text(encoding="utf-8"))
            assert pending_payload["summary"]["status"] == "pending"
            assert pending_payload["summary"]["status_counts"] == {
                "failed": 0,
                "pending": 1,
                "succeeded": 1,
            }
            assert pending_payload["pages"][1]["source"] == "pending"
            assert pending_payload["pages"][1]["status"] == "pending"
            assert pending_payload["pages"][1]["attempts"] == 0
            raise RuntimeError("temporary tesseract failure api_key=abc123 sk-testSECRET123")
        return "OCR page 2 after retry"

    monkeypatch.setattr(extractors.importlib, "import_module", fake_import_module)
    monkeypatch.setattr("librarian.ingest.extractors._ocr_pdf_page", flaky_ocr_pdf_page)
    first = PdfExtractor(ocr_correction_mode="never")
    first.set_page_manifest_path(manifest)

    with pytest.raises(RuntimeError, match="Unable to OCR scanned PDF page 2") as exc_info:
        await first.extract(path)

    error_text = str(exc_info.value)
    assert "api_key=[REDACTED]" in error_text
    assert "[REDACTED]" in error_text
    assert "abc123" not in error_text
    assert "sk-testSECRET123" not in error_text
    failed_payload = json.loads(manifest.read_text(encoding="utf-8"))
    assert failed_payload["pages"][1]["status"] == "failed"
    assert (
        failed_payload["pages"][1]["error"]
        == "temporary tesseract failure api_key=[REDACTED] [REDACTED]"
    )
    assert "abc123" not in failed_payload["pages"][1]["error"]
    assert "sk-testSECRET123" not in failed_payload["pages"][1]["error"]
    assert failed_payload["summary"]["status"] == "failed"
    assert failed_payload["summary"]["status_counts"] == {
        "failed": 1,
        "pending": 0,
        "succeeded": 1,
    }
    assert failed_payload["summary"]["warning_counts"] == {"ocr-page-failed": 1}
    assert failed_payload["pages"][1]["attempts"] == 1
    assert isinstance(failed_payload["pages"][1]["duration_ms"], float)

    second = PdfExtractor(ocr_correction_mode="never")
    second.set_page_manifest_path(manifest)

    text = await second.extract(path)

    retried_payload = json.loads(manifest.read_text(encoding="utf-8"))
    assert calls == 2
    assert "OCR page 2 after retry" in text
    assert retried_payload["summary"]["status"] == "succeeded"
    assert retried_payload["summary"]["attempts"] == 2
    assert retried_payload["pages"][1]["status"] == "succeeded"
    assert retried_payload["pages"][1]["attempts"] == 2
    assert retried_payload["pages"][1]["text"] == "OCR page 2 after retry"


@pytest.mark.asyncio
async def test_pdf_extractor_rejects_symlink_page_manifest(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    path = tmp_path / "fixture.pdf"
    path.write_bytes(b"%PDF")
    outside = tmp_path / "outside-pages.json"
    outside.write_text("keep", encoding="utf-8")
    manifest = tmp_path / "fixture.md.pages.json"
    manifest.symlink_to(outside)

    class FakePage:
        @staticmethod
        def extract_text() -> str:
            return "Text page 1"

    class FakePdf:
        pages = [FakePage()]

        def __enter__(self) -> "FakePdf":
            return self

        def __exit__(self, *args: object) -> None:
            del args

    class FakePdfPlumber:
        @staticmethod
        def open(path: Path) -> FakePdf:
            del path
            return FakePdf()

    def fake_import_module(name: str) -> object:
        if name == "pdfplumber":
            return FakePdfPlumber
        return __import__(name)

    monkeypatch.setattr(extractors.importlib, "import_module", fake_import_module)
    extractor = PdfExtractor(ocr_correction_mode="never")
    extractor.set_page_manifest_path(manifest)

    with pytest.raises(ValueError, match="manifest path must not be a symlink"):
        await extractor.extract(path)

    assert outside.read_text(encoding="utf-8") == "keep"


@pytest.mark.asyncio
async def test_pdf_page_manifest_resume_rejects_oversized_manifest(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    path = tmp_path / "fixture.pdf"
    path.write_bytes(b"%PDF")
    manifest = tmp_path / "fixture.md.pages.json"
    manifest.write_text('{"pages":[]}', encoding="utf-8")

    class FakePage:
        @staticmethod
        def extract_text() -> str:
            return "Text page 1"

    class FakePdf:
        pages = [FakePage()]

        def __enter__(self) -> "FakePdf":
            return self

        def __exit__(self, *args: object) -> None:
            del args

    class FakePdfPlumber:
        @staticmethod
        def open(path: Path) -> FakePdf:
            del path
            return FakePdf()

    def fake_import_module(name: str) -> object:
        if name == "pdfplumber":
            return FakePdfPlumber
        return __import__(name)

    monkeypatch.setattr(extractors.importlib, "import_module", fake_import_module)
    monkeypatch.setattr(extractors, "_MAX_PDF_PAGE_MANIFEST_BYTES", 4)
    extractor = PdfExtractor(ocr_correction_mode="never")
    extractor.set_page_manifest_path(manifest)

    with pytest.raises(ValueError, match="PDF page manifest exceeds"):
        await extractor.extract(path)


@pytest.mark.asyncio
async def test_pdf_extractor_fails_when_mixed_page_ocr_fails(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    path = tmp_path / "fixture.pdf"
    path.write_bytes(b"%PDF")

    class FakePage:
        def __init__(self, text: str | None) -> None:
            self._text = text

        def extract_text(self) -> str | None:
            return self._text

    class FakePdf:
        pages = [FakePage("Text page"), FakePage(None)]

        def __enter__(self) -> "FakePdf":
            return self

        def __exit__(self, *args: object) -> None:
            del args

    class FakePdfPlumber:
        @staticmethod
        def open(path: Path) -> FakePdf:
            del path
            return FakePdf()

    def fake_import_module(name: str) -> object:
        if name == "pdfplumber":
            return FakePdfPlumber
        return __import__(name)

    def fake_ocr_pdf_page(*args: object, **kwargs: object) -> str:
        del args, kwargs
        raise RuntimeError("missing tesseract")

    monkeypatch.setattr(extractors.importlib, "import_module", fake_import_module)
    monkeypatch.setattr("librarian.ingest.extractors._ocr_pdf_page", fake_ocr_pdf_page)

    with pytest.raises(RuntimeError, match="Unable to OCR scanned PDF page 2"):
        await PdfExtractor(ocr_correction_mode="never").extract(path)


@pytest.mark.asyncio
async def test_pdf_extractor_fails_when_scanned_pages_exceed_ocr_limit(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    path = tmp_path / "fixture.pdf"
    path.write_bytes(b"%PDF")

    class FakePage:
        def __init__(self, text: str | None) -> None:
            self._text = text

        def extract_text(self) -> str | None:
            return self._text

    class FakePdf:
        pages = [FakePage("Text page"), FakePage(None), FakePage(None)]

        def __enter__(self) -> "FakePdf":
            return self

        def __exit__(self, *args: object) -> None:
            del args

    class FakePdfPlumber:
        @staticmethod
        def open(path: Path) -> FakePdf:
            del path
            return FakePdf()

    def fake_import_module(name: str) -> object:
        if name == "pdfplumber":
            return FakePdfPlumber
        return __import__(name)

    def fake_ocr_pdf_page(*args: object, **kwargs: object) -> str:
        del args, kwargs
        return "OCR page 2"

    monkeypatch.setattr(extractors.importlib, "import_module", fake_import_module)
    monkeypatch.setattr("librarian.ingest.extractors._ocr_pdf_page", fake_ocr_pdf_page)

    with pytest.raises(ValueError, match="exceeding OCR page limit 1"):
        await PdfExtractor(ocr_pdf_max_pages=1).extract(path)


@pytest.mark.asyncio
async def test_pdf_extractor_fails_when_page_count_exceeds_limit(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    path = tmp_path / "fixture.pdf"
    path.write_bytes(b"%PDF")

    class FakePage:
        @staticmethod
        def extract_text() -> str:
            return "page text"

    class FakePdf:
        pages = [FakePage(), FakePage()]

        def __enter__(self) -> "FakePdf":
            return self

        def __exit__(self, *args: object) -> None:
            del args

    class FakePdfPlumber:
        @staticmethod
        def open(path: Path) -> FakePdf:
            del path
            return FakePdf()

    def fake_import_module(name: str) -> object:
        if name == "pdfplumber":
            return FakePdfPlumber
        return __import__(name)

    monkeypatch.setattr(extractors.importlib, "import_module", fake_import_module)

    with pytest.raises(ValueError, match="2 pages, exceeding configured limit 1"):
        await PdfExtractor(max_pages=1).extract(path)


@pytest.mark.asyncio
async def test_pdf_extractor_llm_corrects_ocr_pages(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    path = tmp_path / "fixture.pdf"
    path.write_bytes(b"%PDF")

    class FakeProvider:
        name = "fake"

        async def complete(self, **kwargs: object) -> str:
            assert "OCR raw" in str(kwargs["user_prompt"])
            return "OCR corrected"

    class FakePage:
        @staticmethod
        def extract_text() -> None:
            return None

    class FakePdf:
        pages = [FakePage()]

        def __enter__(self) -> "FakePdf":
            return self

        def __exit__(self, *args: object) -> None:
            del args

    class FakePdfPlumber:
        @staticmethod
        def open(path: Path) -> FakePdf:
            del path
            return FakePdf()

    def fake_import_module(name: str) -> object:
        if name == "pdfplumber":
            return FakePdfPlumber
        return __import__(name)

    def fake_ocr_pdf_page(*args: object, **kwargs: object) -> str:
        del args, kwargs
        return "OCR raw"

    monkeypatch.setattr(extractors.importlib, "import_module", fake_import_module)
    monkeypatch.setattr("librarian.ingest.extractors._ocr_pdf_page", fake_ocr_pdf_page)

    extractor = PdfExtractor(ocr_correction_provider=FakeProvider())
    text = await extractor.extract(path)

    assert "OCR corrected" in text
    assert "corrected: true" in text
    assert extractor.last_metadata is not None
    pages = extractor.last_metadata["pages"]
    assert isinstance(pages, list)
    assert pages[0]["corrected"] is True


@pytest.mark.asyncio
async def test_pdf_page_manifest_preserves_raw_and_corrected_ocr_text(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    path = tmp_path / "fixture.pdf"
    manifest = tmp_path / "fixture.md.pages.json"
    path.write_bytes(b"%PDF")

    class FakeProvider:
        name = "fake"

        async def complete(self, **kwargs: object) -> str:
            assert "Sadd1e" in str(kwargs["user_prompt"])
            return "Saddle fit and canter transitions."

    class FakePage:
        @staticmethod
        def extract_text() -> None:
            return None

    class FakePdf:
        pages = [FakePage()]

        def __enter__(self) -> "FakePdf":
            return self

        def __exit__(self, *args: object) -> None:
            del args

    class FakePdfPlumber:
        @staticmethod
        def open(path: Path) -> FakePdf:
            del path
            return FakePdf()

    def fake_import_module(name: str) -> object:
        if name == "pdfplumber":
            return FakePdfPlumber
        return __import__(name)

    def fake_ocr_pdf_page(*args: object, **kwargs: object) -> OcrTextResult:
        del args, kwargs
        return OcrTextResult(text="Sadd1e fit and canter transit10ns.", confidence=42.0)

    monkeypatch.setattr(extractors.importlib, "import_module", fake_import_module)
    monkeypatch.setattr("librarian.ingest.extractors._ocr_pdf_page", fake_ocr_pdf_page)
    extractor = PdfExtractor(ocr_correction_provider=FakeProvider())
    extractor.set_page_manifest_path(manifest)

    await extractor.extract(path)

    payload = json.loads(manifest.read_text(encoding="utf-8"))
    page = payload["pages"][0]
    assert page["raw_text"] == "Sadd1e fit and canter transit10ns."
    assert page["corrected_text"] == "Saddle fit and canter transitions."
    assert page["text"] == "Saddle fit and canter transitions."
    assert page["raw_chars"] == len("Sadd1e fit and canter transit10ns.")


@pytest.mark.asyncio
async def test_pdf_page_manifest_records_ocr_correction_quality_warnings(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    path = tmp_path / "fixture.pdf"
    manifest = tmp_path / "fixture.md.pages.json"
    path.write_bytes(b"%PDF")
    repeated_tail = " The correction tail repeats without new evidence." * 8

    class FakeProvider:
        name = "fake"

        async def complete(self, **kwargs: object) -> str:
            del kwargs
            return f"OCR corrected intro.{repeated_tail}"

    class FakePage:
        @staticmethod
        def extract_text() -> None:
            return None

    class FakePdf:
        pages = [FakePage()]

        def __enter__(self) -> "FakePdf":
            return self

        def __exit__(self, *args: object) -> None:
            del args

    class FakePdfPlumber:
        @staticmethod
        def open(path: Path) -> FakePdf:
            del path
            return FakePdf()

    def fake_import_module(name: str) -> object:
        if name == "pdfplumber":
            return FakePdfPlumber
        return __import__(name)

    def fake_ocr_pdf_page(*args: object, **kwargs: object) -> OcrTextResult:
        del args, kwargs
        return OcrTextResult(text="OCR raw", confidence=42.0)

    monkeypatch.setattr(extractors.importlib, "import_module", fake_import_module)
    monkeypatch.setattr("librarian.ingest.extractors._ocr_pdf_page", fake_ocr_pdf_page)
    extractor = PdfExtractor(ocr_correction_provider=FakeProvider())
    extractor.set_page_manifest_path(manifest)

    await extractor.extract(path)

    payload = json.loads(manifest.read_text(encoding="utf-8"))
    page = payload["pages"][0]
    assert "repeated-tail" in page["warnings"]
    assert page["corrected"] is True
    assert extractor.last_metadata is not None
    metadata_pages = extractor.last_metadata["pages"]
    assert isinstance(metadata_pages, list)
    assert "repeated-tail" in metadata_pages[0]["warnings"]


@pytest.mark.asyncio
async def test_pdf_page_manifest_records_preserved_page_image_path(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    path = tmp_path / "fixture.pdf"
    manifest = tmp_path / "fixture.md.pages.json"
    path.write_bytes(b"%PDF")

    class FakePage:
        @staticmethod
        def extract_text() -> None:
            return None

    class FakePdf:
        pages = [FakePage()]

        def __enter__(self) -> "FakePdf":
            return self

        def __exit__(self, *args: object) -> None:
            del args

    class FakePdfPlumber:
        @staticmethod
        def open(path: Path) -> FakePdf:
            del path
            return FakePdf()

    def fake_import_module(name: str) -> object:
        if name == "pdfplumber":
            return FakePdfPlumber
        return __import__(name)

    def fake_ocr_pdf_page(*args: object, **kwargs: object) -> OcrTextResult:
        del args
        image_path = kwargs["image_artifact_path"]
        assert image_path == tmp_path / "fixture.md.pages.page-0001.png"
        return OcrTextResult(text="OCR raw", confidence=90.0, image_path=str(image_path))

    monkeypatch.setattr(extractors.importlib, "import_module", fake_import_module)
    monkeypatch.setattr("librarian.ingest.extractors._ocr_pdf_page", fake_ocr_pdf_page)

    extractor = PdfExtractor(
        ocr_correction_mode="never",
        ocr_preserve_page_images=True,
    )
    extractor.set_page_manifest_path(manifest)
    await extractor.extract(path)

    payload = json.loads(manifest.read_text(encoding="utf-8"))
    page = payload["pages"][0]
    assert page["image_path"] == str(tmp_path / "fixture.md.pages.page-0001.png")
    assert extractor.last_metadata is not None
    metadata_pages = extractor.last_metadata["pages"]
    assert isinstance(metadata_pages, list)
    metadata_page = cast(dict[str, object], metadata_pages[0])
    assert isinstance(metadata_page, dict)
    assert metadata_page["image_path"] == str(tmp_path / "fixture.md.pages.page-0001.png")


@pytest.mark.asyncio
async def test_pdf_extractor_rejects_oversized_ocr_correction(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    path = tmp_path / "fixture.pdf"
    path.write_bytes(b"%PDF")

    class FakeProvider:
        name = "fake"

        async def complete(self, **kwargs: object) -> str:
            del kwargs
            return "x" * 11

    class FakePage:
        @staticmethod
        def extract_text() -> None:
            return None

    class FakePdf:
        pages = [FakePage()]

        def __enter__(self) -> "FakePdf":
            return self

        def __exit__(self, *args: object) -> None:
            del args

    class FakePdfPlumber:
        @staticmethod
        def open(path: Path) -> FakePdf:
            del path
            return FakePdf()

    def fake_import_module(name: str) -> object:
        if name == "pdfplumber":
            return FakePdfPlumber
        return __import__(name)

    def fake_ocr_pdf_page(*args: object, **kwargs: object) -> OcrTextResult:
        del args, kwargs
        return OcrTextResult(text="OCR raw", confidence=42.0)

    monkeypatch.setattr(extractors.importlib, "import_module", fake_import_module)
    monkeypatch.setattr("librarian.ingest.extractors._ocr_pdf_page", fake_ocr_pdf_page)
    extractor = PdfExtractor(
        ocr_correction_provider=FakeProvider(),
        ocr_max_correction_response_chars=10,
    )

    with pytest.raises(RuntimeError, match="exceeding configured limit 10"):
        await extractor.extract(path)


@pytest.mark.asyncio
async def test_pdf_extractor_low_confidence_correction_uses_tesseract_confidence(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    path = tmp_path / "fixture.pdf"
    path.write_bytes(b"%PDF")
    calls = 0

    class FakeProvider:
        name = "fake"

        async def complete(self, **kwargs: object) -> str:
            del kwargs
            nonlocal calls
            calls += 1
            return "OCR corrected"

    class FakePage:
        @staticmethod
        def extract_text() -> None:
            return None

    class FakePdf:
        pages = [FakePage()]

        def __enter__(self) -> "FakePdf":
            return self

        def __exit__(self, *args: object) -> None:
            del args

    class FakePdfPlumber:
        @staticmethod
        def open(path: Path) -> FakePdf:
            del path
            return FakePdf()

    def fake_import_module(name: str) -> object:
        if name == "pdfplumber":
            return FakePdfPlumber
        return __import__(name)

    def fake_ocr_pdf_page(*args: object, **kwargs: object) -> OcrTextResult:
        del args, kwargs
        return OcrTextResult(text="OCR raw", confidence=62.5)

    monkeypatch.setattr(extractors.importlib, "import_module", fake_import_module)
    monkeypatch.setattr("librarian.ingest.extractors._ocr_pdf_page", fake_ocr_pdf_page)

    extractor = PdfExtractor(
        ocr_correction_provider=FakeProvider(),
        ocr_correction_mode="low-confidence",
        ocr_low_confidence_threshold=85,
    )
    text = await extractor.extract(path)

    assert calls == 1
    assert "OCR corrected" in text
    assert extractor.last_metadata is not None
    pages = extractor.last_metadata["pages"]
    assert isinstance(pages, list)
    assert pages[0]["confidence"] == 62.5
    assert pages[0]["corrected"] is True
    assert pages[0]["warnings"] == ["low-ocr-confidence"]


@pytest.mark.asyncio
async def test_pdf_page_manifest_records_missing_confidence_warning(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    path = tmp_path / "fixture.pdf"
    manifest = tmp_path / "fixture.md.pages.json"
    path.write_bytes(b"%PDF")

    class FakePage:
        @staticmethod
        def extract_text() -> None:
            return None

    class FakePdf:
        pages = [FakePage()]

        def __enter__(self) -> "FakePdf":
            return self

        def __exit__(self, *args: object) -> None:
            del args

    class FakePdfPlumber:
        @staticmethod
        def open(path: Path) -> FakePdf:
            del path
            return FakePdf()

    def fake_import_module(name: str) -> object:
        if name == "pdfplumber":
            return FakePdfPlumber
        return __import__(name)

    def fake_ocr_pdf_page(*args: object, **kwargs: object) -> OcrTextResult:
        del args, kwargs
        return OcrTextResult(text="OCR raw", confidence=None)

    monkeypatch.setattr(extractors.importlib, "import_module", fake_import_module)
    monkeypatch.setattr("librarian.ingest.extractors._ocr_pdf_page", fake_ocr_pdf_page)
    extractor = PdfExtractor(ocr_correction_mode="never")
    extractor.set_page_manifest_path(manifest)

    await extractor.extract(path)

    payload = json.loads(manifest.read_text(encoding="utf-8"))
    page = payload["pages"][0]
    assert page["confidence"] is None
    assert page["warnings"] == ["missing-ocr-confidence"]


@pytest.mark.asyncio
async def test_pdf_page_manifest_records_ocr_rotation_retry(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    path = tmp_path / "fixture.pdf"
    manifest = tmp_path / "fixture.md.pages.json"
    path.write_bytes(b"%PDF")

    class FakePage:
        @staticmethod
        def extract_text() -> None:
            return None

    class FakePdf:
        pages = [FakePage()]

        def __enter__(self) -> "FakePdf":
            return self

        def __exit__(self, *args: object) -> None:
            del args

    class FakePdfPlumber:
        @staticmethod
        def open(path: Path) -> FakePdf:
            del path
            return FakePdf()

    def fake_import_module(name: str) -> object:
        if name == "pdfplumber":
            return FakePdfPlumber
        return __import__(name)

    def fake_ocr_pdf_page(*args: object, **kwargs: object) -> OcrTextResult:
        assert kwargs["rotation_retry"] is True
        return OcrTextResult(text="OCR rotated", confidence=92.0, rotation_degrees=90)

    monkeypatch.setattr(extractors.importlib, "import_module", fake_import_module)
    monkeypatch.setattr("librarian.ingest.extractors._ocr_pdf_page", fake_ocr_pdf_page)
    extractor = PdfExtractor(ocr_correction_mode="never", ocr_rotation_retry=True)
    extractor.set_page_manifest_path(manifest)

    await extractor.extract(path)

    payload = json.loads(manifest.read_text(encoding="utf-8"))
    page = payload["pages"][0]
    assert payload["extraction_config"]["ocr_rotation_retry"] is True
    assert page["rotation_degrees"] == 90
    assert page["warnings"] == ["ocr-rotation-retry"]
    assert extractor.last_metadata is not None
    pages = extractor.last_metadata["pages"]
    assert isinstance(pages, list)
    assert pages[0]["rotation_degrees"] == 90


@pytest.mark.asyncio
async def test_pdf_extractor_low_confidence_skips_high_confidence_pages(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    path = tmp_path / "fixture.pdf"
    path.write_bytes(b"%PDF")

    class FakeProvider:
        name = "fake"

        async def complete(self, **kwargs: object) -> str:
            del kwargs
            raise AssertionError("high-confidence OCR should not be corrected")

    class FakePage:
        @staticmethod
        def extract_text() -> None:
            return None

    class FakePdf:
        pages = [FakePage()]

        def __enter__(self) -> "FakePdf":
            return self

        def __exit__(self, *args: object) -> None:
            del args

    class FakePdfPlumber:
        @staticmethod
        def open(path: Path) -> FakePdf:
            del path
            return FakePdf()

    def fake_import_module(name: str) -> object:
        if name == "pdfplumber":
            return FakePdfPlumber
        return __import__(name)

    def fake_ocr_pdf_page(*args: object, **kwargs: object) -> OcrTextResult:
        del args, kwargs
        return OcrTextResult(text="OCR raw", confidence=96.0)

    monkeypatch.setattr(extractors.importlib, "import_module", fake_import_module)
    monkeypatch.setattr("librarian.ingest.extractors._ocr_pdf_page", fake_ocr_pdf_page)

    extractor = PdfExtractor(
        ocr_correction_provider=FakeProvider(),
        ocr_correction_mode="low-confidence",
        ocr_low_confidence_threshold=85,
    )
    text = await extractor.extract(path)

    assert "OCR raw" in text
    assert "corrected: false" in text


def test_parse_tesseract_tsv_confidence_averages_word_confidence() -> None:
    tsv = "\n".join(
        [
            "level\tpage_num\tblock_num\tpar_num\tline_num\tword_num\tconf\ttext",
            "5\t1\t1\t1\t1\t1\t80.5\tAlpha",
            "5\t1\t1\t1\t1\t2\t-1\t",
            "5\t1\t1\t1\t1\t3\t90.5\tBeta",
        ]
    )

    assert extractors.parse_tesseract_tsv_confidence(tsv) == 85.5


def test_ocr_image_rotation_retry_selects_highest_confidence(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source = tmp_path / "page.png"
    source.write_bytes(b"image")
    calls: list[str] = []

    def fake_ocr_image_result(path: Path, **kwargs: object) -> OcrTextResult:
        del kwargs
        calls.append(path.name)
        if "rotated-90" in path.name:
            return OcrTextResult(text="Rotated text", confidence=96.0)
        if "rotated-180" in path.name:
            return OcrTextResult(text="Upside down", confidence=12.0)
        if "rotated-270" in path.name:
            return OcrTextResult(text="Sideways", confidence=35.0)
        return OcrTextResult(text="Original text", confidence=40.0)

    def fake_rotate_ocr_image(path: Path, *, degrees: int, output_dir: Path) -> Path:
        del path
        rotated = output_dir / f"page.rotated-{degrees}.png"
        rotated.write_bytes(b"rotated")
        return rotated

    monkeypatch.setattr("librarian.ingest.extractors._ocr_image_result", fake_ocr_image_result)
    monkeypatch.setattr("librarian.ingest.extractors._rotate_ocr_image", fake_rotate_ocr_image)

    rotation_retry = cast(Any, extractors)._ocr_image_result_with_rotation_retry
    result = rotation_retry(
        source,
        language="eng",
        timeout_seconds=120,
        preprocess_mode="none",
        threshold=180,
        rotation_retry=True,
        low_confidence_threshold=85,
        output_dir=tmp_path,
    )

    assert result.text == "Rotated text"
    assert result.confidence == 96.0
    assert result.rotation_degrees == 90
    assert calls == [
        "page.png",
        "page.rotated-90.png",
        "page.rotated-180.png",
        "page.rotated-270.png",
    ]


@pytest.mark.asyncio
async def test_image_ocr_applies_threshold_preprocessing(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    image_module = pytest.importorskip("PIL.Image")
    source = tmp_path / "scan.png"
    image = image_module.new("L", (2, 1))
    image.putdata([40, 220])
    image.save(source)

    def fake_run_tesseract(path: Path, **kwargs: object) -> subprocess.CompletedProcess[str]:
        del kwargs
        output = image_module.open(path).convert("L")
        assert list(output.tobytes()) == [0, 255]
        return subprocess.CompletedProcess(args=["tesseract"], returncode=0, stdout="OCR text\n")

    def fake_which(name: str) -> str:
        return name

    monkeypatch.setattr(extractors.shutil, "which", fake_which)
    monkeypatch.setattr("librarian.ingest.extractors._run_tesseract", fake_run_tesseract)

    text = await ImageOcrExtractor(preprocess_mode="threshold", threshold=180).extract(source)

    assert text == "OCR text"


@pytest.mark.asyncio
async def test_pdf_extractor_passes_ocr_preprocessing_options(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    path = tmp_path / "fixture.pdf"
    path.write_bytes(b"%PDF")
    seen: dict[str, object] = {}

    class FakePage:
        @staticmethod
        def extract_text() -> None:
            return None

    class FakePdf:
        pages = [FakePage()]

        def __enter__(self) -> "FakePdf":
            return self

        def __exit__(self, *args: object) -> None:
            del args

    class FakePdfPlumber:
        @staticmethod
        def open(path: Path) -> FakePdf:
            del path
            return FakePdf()

    def fake_import_module(name: str) -> object:
        if name == "pdfplumber":
            return FakePdfPlumber
        return __import__(name)

    def fake_ocr_pdf_page(*args: object, **kwargs: object) -> OcrTextResult:
        del args
        seen.update(kwargs)
        return OcrTextResult(text="OCR raw", confidence=92.0)

    monkeypatch.setattr(extractors.importlib, "import_module", fake_import_module)
    monkeypatch.setattr("librarian.ingest.extractors._ocr_pdf_page", fake_ocr_pdf_page)

    extractor = PdfExtractor(
        ocr_correction_mode="never",
        ocr_preprocess_mode="deskew",
        ocr_threshold=155,
    )
    await extractor.extract(path)

    assert seen["preprocess_mode"] == "deskew"
    assert seen["threshold"] == 155
    assert extractor.last_metadata is not None
    assert extractor.last_metadata["extraction_config"] == {
        "ocr_language": "eng",
        "ocr_timeout_seconds": 120,
        "ocr_pdf_dpi": 200,
        "ocr_preprocess_mode": "deskew",
        "ocr_threshold": 155,
        "ocr_preserve_page_images": False,
        "ocr_rotation_retry": False,
        "ocr_correction_mode": "never",
        "ocr_correction_model": "mock-cleaner",
        "ocr_low_confidence_threshold": 85.0,
        "ocr_max_correction_response_chars": 2 * 1024 * 1024,
    }


@pytest.mark.asyncio
async def test_pdf_extractor_defaults_to_1000_ocr_pages() -> None:
    extractor = PdfExtractor()

    assert extractor.ocr_pdf_max_pages == 1_000


@pytest.mark.asyncio
async def test_mock_provider_ocr_correction_returns_only_page_text() -> None:
    provider = MockLLMProvider()

    text = await provider.complete(
        system_prompt="correct OCR",
        user_prompt="Correct OCR text from PDF page 1. Return only corrected text.\n\nOCR raw text",
        model="mock-cleaner",
        max_tokens=128,
        temperature=0.0,
    )

    assert text == "OCR raw text"
